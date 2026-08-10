from __future__ import annotations

import argparse
import json
from collections import Counter, OrderedDict
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# Design system
# Base preset: compact_reference_guide.
# Named overrides used consistently:
# - FusionDigital CJK typography: Microsoft YaHei replaces Calibri.
# - FusionDigital palette: deep green + orange + cyan replaces default blue.
# - Editorial report cover: editorial_cover header pattern.
PAGE_WIDTH_DXA = 12240
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
FONT = "Microsoft YaHei"
INK = "102018"
MUTED = "5E6B64"
DEEP = "146B55"
DEEP_DARK = "0D2B24"
FUSION = "FF8738"
CYAN = "65E6D2"
PURPLE = "8C78D9"
PALE = "E8F0EB"
PALE_ORANGE = "FFF0E5"
PALE_PURPLE = "F0EDFC"
LINE = "C8D2CB"
WHITE = "FFFFFF"

DOMAIN_ORDER = ["physics", "engineering", "control", "diagnostics", "energy", "auxiliary", "data", "hmi", "integration"]
DOMAIN_META = {
    "physics": ("物理模拟", "PHYSICS", "domain-physics-dark-image2.png"),
    "engineering": ("工程仿真", "ENGINEERING", "domain-engineering-dark-image2.png"),
    "control": ("集成控制", "INTEGRATED CONTROL", "domain-integrated-control-dark-image2.png"),
    "diagnostics": ("智能诊断", "INTELLIGENT DIAGNOSTICS", "domain-intelligent-diagnostics-dark-image2.png"),
    "energy": ("能量转化", "ENERGY CONVERSION", "domain-energy-conversion-dark-image2.png"),
    "auxiliary": ("辅机模拟", "AUXILIARY SYSTEMS", "domain-auxiliary-systems-dark-image2.png"),
    "data": ("数据基座", "DATA FOUNDATION", "domain-data-foundation-dark-image2.png"),
    "hmi": ("人机交互", "HUMAN–MACHINE INTERACTION", "domain-human-machine-interaction-dark-image2.png"),
    "integration": ("总体集成", "WHOLE-PLANT INTEGRATION", "domain-whole-plant-integration-dark-image2.png"),
}
DOMAIN_SYNTHESIS = {
    "physics": {
        "need": "把输运、湍流、平衡、边界和材料相互作用等慢模型转化为可用于场景优化、状态估计和控制设计的快速预测能力。",
        "state": "代理输运、可微核心输运和神经状态空间最成熟；典型路径是用 QuaLiKiz、高保真模拟或装置数据训练快速模型，再嵌入集成模拟。",
        "gap": "代理继承父模型偏差，跨尺度和核心—边界耦合仍弱；训练域外推、守恒约束与不确定度传播是进入数字孪生的关键门槛。",
    },
    "engineering": {
        "need": "降低电磁、结构、热流体、中子和材料计算在参数扫描、逆设计、实时估计与维护决策中的计算代价。",
        "state": "概念设计、几何优化和多物理降阶已有论文级原型；真实工程传感器闭环证据明显少于等离子体控制与诊断。",
        "gap": "制造公差、材料退化、辐照、局部热点和复杂接触等效应难以覆盖；商业求解器数据与公开代码之间存在较大复现鸿沟。",
    },
    "control": {
        "need": "在共享执行器、强耦合和不稳定边界下，实时协调磁位形、剖面、加热、密度与 MHD 稳定性。",
        "state": "TCV、DIII-D、KSTAR 等装置已给出真实闭环证据，是智能原生中实验成熟度最高的方向。",
        "gap": "策略通常针对有限工况；多目标冲突、最坏时延、执行器故障、域外检测和确定性安全包络尚未达到电厂级。",
    },
    "diagnostics": {
        "need": "从高维、异步、噪声与缺失的传感器中形成可信状态，识别破裂、不稳定性、异常和部件健康。",
        "state": "破裂预测、虚拟诊断、断层重建和多模态预训练快速发展；FusionMAE、TokaMind 等探索统一状态表征。",
        "gap": "跨装置标定、真实故障样本、长期漂移、误报警代价和辐照环境退化尚未被系统验证。",
    },
    "energy": {
        "need": "把包层取热、一次/二次回路、功率循环、储能与并网约束纳入快速设计和运行优化。",
        "state": "直接面向聚变的 AI 研究较少，方法多从核电、热工系统和系统工程迁移，现阶段以概念设计和离线优化为主。",
        "gap": "聚变热源的脉冲/瞬态、氚系统、热循环疲劳和电网需求尚无成体系的运行数据闭环。",
    },
    "auxiliary": {
        "need": "提升加热与电流驱动、低温、真空、燃料/氚、冷却、电源、遥操作和维护系统的效率与可用率。",
        "state": "RF/NBI 快速代理、设备异常检测和机器人感知是最现实的切入点；多数证据仍是子系统或邻近工业场景。",
        "gap": "数据孤岛、商业设备接口、故障标签稀缺和跨系统失效传播使全厂协同困难。",
    },
    "data": {
        "need": "统一时间、单位、配置、标识、本体、质量、权限、版本和来源，使模型、实验与仿真结果可复现、可追责。",
        "state": "MDSplus、UDA、IMAS/OMAS、Fusion Data Platform、FAIR-MAST、TokSearch 等形成分层基础设施。",
        "gap": "跨装置语义和权限仍碎片化；开放数据、诊断元数据、稳定基准和训练切分远不足以支撑通用基础模型。",
    },
    "hmi": {
        "need": "帮助研究者与操作者通过自然语言检索、解释、规划、分析与协同，但不模糊责任和权限。",
        "state": "聚变领域 LLM、RAG、数值数据智能体和运行知识助手已出现，主要处于原型或演示阶段。",
        "gap": "引用核验、程序约束、人因验证、保密、权限、工具调用可控性和专家责任尚未形成工程标准。",
    },
    "integration": {
        "need": "把物理、工程、控制、诊断、辅机、经济和安全目标组织为多保真、可追溯、可迭代的全厂决策闭环。",
        "state": "FUSE、PROCESS/FAROES、可微模拟和优化框架已能支撑设计空间探索，但大多仍是离线集成模拟。",
        "gap": "距离数字孪生仍缺实时权威状态、在线校准、全厂故障传播、生命周期配置、UQ 和安全论证闭环。",
    },
}

EVIDENCE_LABELS = {
    "E0": "概念 / 方法",
    "E1": "仿真或合成数据",
    "E2": "真实装置离线数据",
    "E3": "实时系统 / HIL / 影子",
    "E4": "真实装置闭环实验",
}
CODE_LABELS = {
    "official-direct": "官方对应实现",
    "official-enabling": "官方使能工具",
    "commercial-enabling": "商业使能工具",
    "community-reproduction": "社区复现",
    "not-public": "未公开",
}

DEPLOYMENT_LABELS = {
    "D1": "概念 / 路线",
    "D2": "离线研究原型",
    "D3": "装置数据验证或运行试点",
    "D4": "生产服务、正式工作流或装置常规使用",
    "D5": "安全关键在线系统或电厂级持续运行",
}

PAPER_SOURCE_LABELS = {
    "peer-reviewed": "同行评审论文",
    "journal-paper": "期刊论文",
    "conference-paper": "会议论文",
    "preprint": "预印本",
    "official-report": "官方报告",
    "technical-report": "技术报告",
    "official-page": "机构 / 项目官方页面",
    "official-documentation": "官方文档",
    "thesis": "学位论文",
    "dataset": "数据集说明",
}

ARTIFACT_TYPE_LABELS = {
    "source-code": "源代码",
    "software": "软件",
    "commercial-software": "商业软件",
    "framework": "框架",
    "model-weights": "模型权重",
    "notebook": "Notebook",
    "dataset": "数据集",
    "documentation": "文档",
    "unknown": "资产类型未标注",
}

ACCESS_LABELS = {
    "open-source": "开源",
    "open": "开放访问",
    "public": "公开",
    "proprietary": "专有 / 商业许可",
    "restricted": "受限访问",
    "not-public": "未公开",
    "unknown": "访问方式未标注",
}

DOMAIN_ALIASES = {
    **{domain: domain for domain in DOMAIN_ORDER},
    "物理模拟": "physics",
    "工程仿真": "engineering",
    "集成控制": "control",
    "智能诊断": "diagnostics",
    "能量转化": "energy",
    "辅机模拟": "auxiliary",
    "数据基座": "data",
    "人机交互": "hmi",
    "总体集成": "integration",
}

def as_list(value: Any) -> list[Any]:
    """Return a schema value as a list without splitting strings."""
    if value is None:
        return []
    if isinstance(value, list):
        return value
    if isinstance(value, tuple):
        return list(value)
    return [value]


def normalize_domain(value: Any) -> str:
    raw = str(value or "").strip()
    return DOMAIN_ALIASES.get(raw, DOMAIN_ALIASES.get(raw.lower(), raw.lower()))


def unique_values(values: Iterable[Any]) -> list[Any]:
    result = []
    seen = set()
    for value in values:
        marker = json.dumps(value, ensure_ascii=False, sort_keys=True) if isinstance(value, (dict, list)) else str(value)
        if marker not in seen:
            seen.add(marker)
            result.append(value)
    return result


def label_for(mapping: dict[str, str], value: Any, missing: str = "未标注") -> str:
    raw = str(value or "").strip()
    return mapping.get(raw, raw or missing)


def evidence_label(level: Any) -> str:
    return label_for(EVIDENCE_LABELS, level, "证据等级未标注")


def deployment_label(level: Any) -> str:
    return label_for(DEPLOYMENT_LABELS, level, "未单独判级")


def paper_source_label(source_type: Any) -> str:
    return label_for(PAPER_SOURCE_LABELS, source_type, "来源类型未标注")


def code_status_label(status: Any) -> str:
    return label_for(CODE_LABELS, status, "代码关系未标注")


def artifact_type_label(artifact_type: Any) -> str:
    return label_for(ARTIFACT_TYPE_LABELS, artifact_type, "资产类型未标注")


def access_label(access: Any) -> str:
    return label_for(ACCESS_LABELS, access, "访问方式未标注")


def normalize_entry(raw: dict[str, Any], index: int) -> dict[str, Any]:
    """Normalize both the legacy single-domain schema and the merged multi-domain schema."""
    item = dict(raw)
    primary = normalize_domain(item.get("primaryDomain") or item.get("domain"))
    if primary not in DOMAIN_META:
        primary = "integration"
    related = []
    for value in as_list(item.get("relatedDomains")):
        domain = normalize_domain(value)
        if domain in DOMAIN_META and domain != primary:
            related.append(domain)

    project_id = str(item.get("projectId") or item.get("id") or f"work-{index:03d}")
    item.update(
        {
            "projectId": project_id,
            "id": str(item.get("id") or project_id),
            "primaryDomain": primary,
            "domain": primary,  # Legacy alias retained for older report helpers.
            "relatedDomains": unique_values(related),
            "deploymentLevel": str(item.get("deploymentLevel") or "").strip(),
            "title": str(item.get("title") or "未命名工作"),
            "year": item.get("year") or "未标注",
            "organization": str(item.get("organization") or "未标注"),
            "problem": str(item.get("problem") or "未标注"),
            "approach": str(item.get("approach") or "未标注"),
            "devices": [str(v) for v in as_list(item.get("devices")) if str(v).strip()],
            "evidenceLevel": str(item.get("evidenceLevel") or "").strip(),
            "evidence": str(item.get("evidence") or "未标注"),
            "data": str(item.get("data") or "未标注"),
            "maturity": str(item.get("maturity") or "未标注"),
            "limitations": str(item.get("limitations") or "未标注"),
            "tags": [str(v) for v in as_list(item.get("tags")) if str(v).strip()],
        }
    )

    papers = []
    for paper in as_list(item.get("papers")):
        if not isinstance(paper, dict):
            continue
        normalized = dict(paper)
        normalized.update(
            {
                "title": str(paper.get("title") or "未命名来源"),
                "year": paper.get("year") or item["year"],
                "venue": str(paper.get("venue") or "未标注"),
                "url": str(paper.get("url") or ""),
                "sourceType": str(paper.get("sourceType") or "").strip(),
            }
        )
        papers.append(normalized)
    item["papers"] = papers

    codes = []
    for code in as_list(item.get("code")):
        if not isinstance(code, dict):
            continue
        normalized = dict(code)
        normalized.update(
            {
                "name": str(code.get("name") or "未命名代码 / 工具"),
                "url": str(code.get("url") or ""),
                "status": str(code.get("status") or "not-public"),
                "relationship": str(code.get("relationship") or "未说明与本工作的关系"),
                "artifactType": str(code.get("artifactType") or "").strip(),
                "access": str(code.get("access") or "").strip(),
            }
        )
        codes.append(normalized)
    item["code"] = codes
    return item


def merge_unique_works(entries: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Deduplicate by projectId while conservatively preserving domain/source relationships."""
    works: OrderedDict[str, dict[str, Any]] = OrderedDict()
    for item in entries:
        key = item["projectId"]
        if key not in works:
            works[key] = item
            continue
        current = works[key]
        candidate_domains = [item["primaryDomain"], *item["relatedDomains"]]
        current["relatedDomains"] = unique_values(
            [*current["relatedDomains"], *(d for d in candidate_domains if d != current["primaryDomain"])]
        )
        for field in ("devices", "tags", "papers", "code"):
            current[field] = unique_values([*current[field], *item[field]])
    return list(works.values())


def entry_domains(item: dict[str, Any]) -> list[str]:
    return unique_values([item["primaryDomain"], *item.get("relatedDomains", [])])


def is_related_work(item: dict[str, Any], domain: str) -> bool:
    return domain in item.get("relatedDomains", []) and item["primaryDomain"] != domain


def domain_display(domain: str) -> str:
    return DOMAIN_META.get(domain, (domain or "未标注", "", ""))[0]


def set_run_font(run, size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None, name: str = FONT):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        element = tc_mar.find(qn(f"w:{margin}"))
        if element is None:
            element = OxmlElement(f"w:{margin}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = LINE, size: str = "4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = TABLE_INDENT_DXA):
    assert sum(widths_dxa) == CONTENT_WIDTH_DXA, (widths_dxa, sum(widths_dxa))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
    set_table_borders(table)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str = INK, size: float = 8.5, align=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def add_hyperlink(paragraph, text: str, url: str, color: str = DEEP):
    part = paragraph.part
    relationship_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), FONT)
    r_pr.append(r_fonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, text, end):
        run._r.append(element)
    set_run_font(run, size=8, color=MUTED)


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    for style_name, size, color, before, after in (
        ("Heading 1", 16, DEEP, 18, 10),
        ("Heading 2", 13, DEEP, 14, 7),
        ("Heading 3", 12, DEEP_DARK, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("Caption",):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(8.5)
        style.font.italic = False
        style.font.color.rgb = RGBColor.from_string(MUTED)
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(9)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True


def configure_headers_footers(section):
    first = section.first_page_header
    p = first.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run("Fusion")
    set_run_font(r1, size=9, color=FUSION, bold=True)
    r2 = p.add_run("Digital")
    set_run_font(r2, size=9, color=CYAN, bold=True)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run("FusionDigital  /  聚变智能原生技术图谱")
    set_run_font(run, size=8, color=MUTED, bold=True)
    p_pr = hp._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), LINE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    label = fp.add_run("新奥聚变人工智能团队  ·  ")
    set_run_font(label, size=8, color=MUTED)
    add_page_number(fp)


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None, color: str = INK, after: float = 6):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first, size=10.5, color=DEEP_DARK, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(rest, size=10.5, color=color)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=10.5, color=color)
    return paragraph


def add_bullet(doc: Document, text: str, *, level: int = 0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.375 + level * 0.2)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=10.2, color=INK)
    return p


def add_kicker(doc: Document, text: str, *, page_break_before: bool = False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.page_break_before = page_break_before
    run = p.add_run(text.upper())
    set_run_font(run, size=8.5, color=FUSION, bold=True)
    run.font.all_caps = True
    return p


def add_figure(doc: Document, path: Path, caption: str, *, width: float = 6.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cp = doc.add_paragraph(style="Caption")
    cr = cp.add_run(caption)
    set_run_font(cr, size=8.5, color=MUTED)


def add_callout(doc: Document, title: str, text: str, *, fill: str = PALE, accent: str = DEEP):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [1050, 8310])
    set_cell_shading(table.cell(0, 0), accent)
    set_cell_shading(table.cell(0, 1), fill)
    set_cell_text(table.cell(0, 0), title, bold=True, color=WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(0, 1), text, color=INK, size=9.2)
    table.rows[0].cells[1].paragraphs[0].paragraph_format.line_spacing = 1.2
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)


def add_metric_strip(doc: Document, metrics: list[tuple[str, str]]):
    table = doc.add_table(rows=1, cols=len(metrics))
    widths = [CONTENT_WIDTH_DXA // len(metrics)] * len(metrics)
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    set_table_geometry(table, widths, indent_dxa=0)
    for index, (value, label) in enumerate(metrics):
        cell = table.cell(0, index)
        set_cell_shading(cell, DEEP_DARK)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(value)
        set_run_font(r, size=17, color=CYAN if index % 2 == 0 else FUSION, bold=True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(label)
        set_run_font(r2, size=7.5, color="B9C9C1")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_simple_table(doc: Document, headers: list[str], rows: Iterable[list[str]], widths_dxa: list[int], *, header_fill: str = DEEP_DARK, font_size: float = 8.2):
    rows = list(rows)
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        set_cell_shading(table.cell(0, idx), header_fill)
        set_cell_text(table.cell(0, idx), header, bold=True, color=WHITE, size=8.2)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            set_cell_text(cells[idx], str(value), size=font_size, color=INK)
            if row_index % 2 == 1:
                set_cell_shading(cells[idx], "F4F6F3")
        prevent_row_split(table.rows[-1])
    set_table_geometry(table, widths_dxa)
    return table


def table_citation(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=7.8, color=MUTED, italic=True)


def add_cover(doc: Document, figures_dir: Path, stats: dict[str, Any]):
    add_kicker(doc, "FusionDigital · Fusion AI-native research atlas")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(26)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("聚变智能原生技术图谱")
    set_run_font(r, size=30, color=DEEP_DARK, bold=True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(18)
    r2 = p2.add_run(f"面向 Tokamak、聚变堆与聚变电厂数字孪生的 {stats['unique_total']} 个唯一工作 / 九域关联覆盖")
    set_run_font(r2, size=13.5, color=DEEP, bold=False)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(18)
    r3 = p3.add_run("RESEARCH LANDSCAPE · EVIDENCE MAP · IMPLEMENTATION ROADMAP")
    set_run_font(r3, size=8.5, color=FUSION, bold=True)
    add_figure(
        doc,
        figures_dir / "fusion-twin-ai-native-overview.png",
        "图 0-1  聚变装置—数字孪生—智能体的受控闭环。概念性科学架构图，用于解释系统关系，不代表数值模拟输出。",
        width=5.95,
    )
    add_metric_strip(
        doc,
        [
            (str(stats["unique_total"]), "个唯一工作"),
            ("9", "个关联覆盖知识域"),
            (str(stats["domain_relationships"]), "条工作—领域关联"),
            (str(stats["device_evidence"]), "项 E2+ 证据"),
        ],
    )
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(10)
    p4.paragraph_format.space_after = Pt(0)
    r4 = p4.add_run("新奥聚变人工智能团队  ·  2026 年 8 月 10 日  ·  tianshao1992@gmail.com")
    set_run_font(r4, size=8.5, color=MUTED)


def add_manual_toc(doc: Document):
    heading = doc.add_heading("阅读导航", level=1)
    heading.paragraph_format.page_break_before = True
    add_body(doc, "本报告按“概念—证据—九域图谱—装置采用—数字孪生差距—FusionDigital 路线”组织。九个知识域均包含代表工作、装置、验证、论文、代码、数据、成熟度与局限。", after=10)
    toc_rows = [
        ("01", "执行摘要：智能原生的价值、现状与边界"),
        ("02", "科普：为什么聚变需要 AI，但不能让 AI 取代物理"),
        ("03", "调研方法：证据等级、部署等级、代码关系与适用域"),
        ("04", "九域总览：唯一工作与领域关联双口径"),
        ("05–13", "九域图谱：物理、工程、控制、诊断、能量、辅机、数据、人机、总体集成"),
        ("14", "装置采用与跨装置迁移"),
        ("15", "现有 AI 工作距离数字孪生的差距"),
        ("16", "FusionDigital 分阶段总体路线"),
        ("17", "联合研发与专家协作清单"),
        ("附录", "论文与代码索引、术语与维护说明"),
    ]
    add_simple_table(doc, ["章节", "内容"], [[a, b] for a, b in toc_rows], [1350, 8010], header_fill=DEEP_DARK, font_size=9)
    table_citation(doc, "导航为内容结构索引；Word 中可使用“导航窗格”按标题快速跳转。")


def add_executive_summary(doc: Document, entries: list[dict[str, Any]], stats: dict[str, Any]):
    heading = doc.add_heading("1  执行摘要：智能原生的价值、现状与边界", level=1)
    heading.paragraph_format.page_break_before = True
    add_callout(
        doc,
        "核心判断",
        "聚变智能原生不是把通用大模型接入控制系统，而是让数据、物理模型、实验、工程仿真、控制器和专家决策共享同一套可追溯状态与证据，并用专用 AI 提速、用智能体编排、用确定性安全层约束。",
        fill=PALE_ORANGE,
        accent=FUSION,
    )
    add_body(
        doc,
        f"本次合并图谱以 projectId 去重，共整理 {stats['unique_total']} 个公开可核验的唯一工作，"
        f"覆盖九个 FusionDigital 知识域。由于一个工作可同时服务主域和关联域，共形成 {stats['domain_relationships']} 条工作—领域关联，"
        f"其中 {stats['primary_relationships']} 条为主域归属、{stats['related_relationships']} 条为关联域映射。唯一工作数不能与各领域章节条目数相加混用。"
        f"其中 {stats['device_evidence']} 项使用真实装置数据或达到更高证据，{stats['direct_code']} 项可确认官方对应实现。",
    )
    doc.add_heading("1.1  六个结论", level=2)
    conclusions = [
        "集成控制是当前实验证据最强的 AI 方向。TCV 的磁位形控制、DIII-D/KSTAR 的不稳定性避免等工作已经触及真实装置闭环，但其适用域仍窄，不能直接外推到燃烧等离子体或电厂连续运行。",
        "智能诊断拥有最丰富的装置数据应用。破裂预测、虚拟诊断、断层重建、缺失信号恢复与基础模型正在汇合，但跨装置语义、标定与分布漂移仍是主要瓶颈。",
        "物理与工程代理的直接价值是把‘算不起’变成‘可搜索、可优化、可用于影子预测’。代理速度很快，但永远继承训练数据和父模型的偏差，必须和高保真模型、实验及不确定度共同交付。",
        "能量转化与辅机模拟的聚变专属 AI 证据最薄弱。近期更适合从 RF/NBI 代理、低温/真空/冷却设备健康、热循环优化和遥操作等边界清晰的子系统切入。",
        "基础模型和智能体带来统一接口与流程自动化，但可信度不能由参数量推断。最有价值的近期形态是‘有引用、有工具白名单、无直接控制权限’的研究与操作副驾驶。",
        "数据基座不是后台 IT 项目，而是智能原生的第一性工程。没有时间对齐、配置、单位、本体、数据质量、版本、权限和证据谱系，就不存在可审计的数字孪生闭环。",
    ]
    for conclusion in conclusions:
        add_bullet(doc, conclusion)
    doc.add_heading("1.2  对 FusionDigital 的建议", level=2)
    add_body(doc, "采用‘窄任务价值—影子孪生—跨域协同—受控自治’路线。当前从 DINA、MEQ 切入控制服务是合理起点；应同时建设模型/数据证据账本，并并行推进工程电磁力—热管理代理、诊断状态表征和仿真智能体。所有工作以可重复评测和实验闭环为共同交付物。")


def add_primer(doc: Document, figures_dir: Path):
    doc.add_heading("2  科普：为什么聚变需要 AI，但不能让 AI 取代物理", level=1)
    doc.add_heading("2.1  聚变问题为什么特别难", level=2)
    add_body(doc, "Tokamak 同时包含微秒级等离子体不稳定性、秒到小时级热与结构响应、天到年级材料退化和维护决策。不同子系统使用不同坐标、网格、采样率、软件和证据标准；实验数据昂贵、工况覆盖稀疏，未来装置又不能靠大量破坏性事件来训练模型。")
    add_body(doc, "因此，AI 的价值主要体现在四件事：从传感器形成状态、把慢模型变快、在庞大设计/运行空间中搜索、以及编排跨软件与跨团队流程。AI 不负责定义物理真理，也不应取代独立安全保护。")
    add_figure(
        doc,
        figures_dir / "domain-ai-native-dark-image2.png",
        "图 2-1  智能原生能力层：状态表征、代理预测、搜索优化与智能体编排围绕物理模型和实验证据展开。概念性科学架构图。",
        width=6.35,
    )
    stack_heading = doc.add_heading("2.2  四层技术栈", level=2)
    stack_heading.paragraph_format.page_break_before = True
    rows = [
        ["机器学习", "分类、回归、异常检测、贝叶斯优化", "破裂预警、参数辨识、设备健康", "概率校准与严格数据切分"],
        ["深度学习", "多模态时序、图像、神经算子、RL", "状态估计、快速代理、闭环控制", "传感器失效、OOD、最坏时延"],
        ["基础模型", "自监督预训练、共享表征、少样本适配", "缺失信号、跨任务接口、统一状态向量", "数据谱系、跨装置独立测试"],
        ["智能体", "检索、代码、仿真、优化、报告工具编排", "实验规划、仿真工厂、操作副驾驶", "最小权限、人工批准、全链审计"],
    ]
    add_simple_table(doc, ["技术层", "核心能力", "适合任务", "最低可信要求"], rows, [1500, 2500, 2600, 2760], font_size=8.4)
    table_citation(doc, "表 2-1  四层能力不是替代关系；专用模型与确定性控制器仍是在线闭环主体。")
    doc.add_heading("2.3  智能原生与传统‘AI 加点’的区别", level=2)
    add_callout(doc, "AI 加点", "单独训练一个模型或部署一个聊天助手；输入、版本、配置和证据常与装置主线分离。", fill="F4F6F3", accent=MUTED)
    add_callout(doc, "智能原生", "从架构起点就定义共享状态、时间轴、配置、权限、证据门和回退机制；AI 是可替换组件，数字线程与安全边界是系统资产。", fill=PALE, accent=DEEP)


def add_methodology(doc: Document, stats: dict[str, Any]):
    doc.add_heading("3  调研方法：证据、部署、代码关系与适用域", level=1)
    add_body(doc, "本报告优先采用同行评审论文、arXiv 原始预印本、DOE/IAEA/装置机构官方页面和作者/项目官方仓库。聚变 AI 领域存在‘论文有结果但无代码’、‘有通用仓库但并非论文实现’、‘计划适配未来装置但尚未实验’等常见混淆，因此将证据强度、部署深度与代码关系作为相互独立的字段标注。")
    doc.add_heading("3.1  证据等级：公开材料直接证明了什么", level=2)
    evidence_rows = [
        ["E0", "概念 / 方法", "提出架构、方法或研发计划；不足以判断装置效果"],
        ["E1", "仿真 / 合成数据", "在父模型、高保真模拟、合成数据或概念设计数据上验证"],
        ["E2", "装置离线数据", "使用真实装置历史数据训练、测试或回放"],
        ["E3", "实时 / HIL / 影子", "进入实时系统、硬件在环或影子运行，但不直接闭环驱动装置"],
        ["E4", "装置闭环实验", "在真实装置中闭环影响执行器、实验轨迹或保护动作"],
    ]
    add_simple_table(doc, ["等级", "名称", "解释"], evidence_rows, [900, 2100, 6360], font_size=8.8)
    table_citation(doc, "表 3-1  只按公开来源直接证明的最高等级标注；‘目标用于 ITER/SPARC’不等同于已在目标装置验证。")
    doc.add_heading("3.2  部署等级：能力进入运行体系有多深", level=2)
    deployment_rows = [
        ["D1", "概念 / 路线", "公开材料主要描述需求、方法或研发路线，尚未形成可复核原型"],
        ["D2", "离线研究原型", "形成可执行研究代码、离线算法或一次性分析流程，尚未进入装置运行链"],
        ["D3", "装置验证 / 运行试点", "使用装置数据验证，或进入实时、影子、HIL 与受控运行试点"],
        ["D4", "正式工作流 / 常规使用", "进入生产服务、正式设计流程或装置常规使用，但不据此宣称承担安全关键职责"],
        ["D5", "安全关键 / 电厂持续运行", "承担安全关键在线功能或电厂级持续运行，并有公开证据证明相应批准、治理与全生命周期保障"],
    ]
    add_simple_table(doc, ["等级", "名称", "解释"], deployment_rows, [900, 2600, 5860], font_size=8.5)
    table_citation(doc, "表 3-2  部署等级不从 evidenceLevel 推导；旧条目没有 deploymentLevel 时明确标为‘未单独判级’。")
    add_callout(
        doc,
        "两个轴不可互推",
        "证据等级回答‘论文或官方材料直接证明到哪一步’，部署等级回答‘该能力在组织、流程、权限和安全治理中处于什么位置’。例如，真实装置离线数据可以形成 E2 证据，但仍可能只是 D2 离线研究原型；正式数据服务可达 D4，却不因此自动成为 E4 装置闭环实验，更不代表达到 D5 安全关键部署。",
        fill=PALE_PURPLE,
        accent=PURPLE,
    )
    doc.add_heading("3.3  代码与资产关系", level=2)
    code_rows = [
        ["官方对应实现", "论文作者或项目方公开，与该工作直接对应的训练/推理代码或权重"],
        ["官方使能工具", "官方仓库可支撑流程或提供父模型，但不是论文模型完整对应实现"],
        ["商业使能工具", "商业或专有软件支撑该工作；记录资产类型和访问方式，但不等同于可公开复现"],
        ["社区复现", "第三方复现或相近实现；只作为学习入口，不视为原始代码"],
        ["未公开", "未发现可确认的对应仓库，或仅有闭源、内部/商业实现"],
    ]
    add_simple_table(doc, ["标记", "判断规则"], code_rows, [2200, 7160], font_size=8.8)
    table_citation(doc, f"表 3-3  当前图谱代码关系计数：{json.dumps(stats['code_counts'], ensure_ascii=False)}。一个工作可能关联多个代码资产。")
    doc.add_heading("3.4  阅读与计数限制", level=2)
    for item in (
        "图谱面向 Tokamak 和聚变电厂数字孪生规划，不声称穷尽所有磁约束、惯性约束或商业内部工作。",
        "商业软件、装置内网和未公开训练数据会造成可复现性盲区；报告明确记录这些缺口。",
        f"唯一工作数按 projectId 去重；当前合并数据为 {stats['unique_total']} 个。领域关联数则统计主域和 relatedDomains，不能与唯一工作数混为一谈。",
        "九域章节同时展示主域工作和关联工作；关联命中会显式标注‘关联工作’，同一项目可能在多个章节出现。",
        "装置适配包含直接验证、训练数据来源和作者提出的未来目标；三者在条目中分别说明。",
    ):
        add_bullet(doc, item)


def add_landscape_overview(doc: Document, entries: list[dict[str, Any]], figures_dir: Path, stats: dict[str, Any]):
    doc.add_heading("4  九域总览：智能原生如何贯穿聚变数字孪生", level=1)
    add_body(doc, f"总览采用双口径：{stats['unique_total']} 个 projectId 去重后的唯一工作，以及 {stats['domain_relationships']} 条工作—领域关联。下表每一行统计该领域的主域命中与关联域命中，因此九行关系数之和大于唯一工作数是正常现象。")
    add_figure(
        doc,
        figures_dir / "integrated-twin-reference-architecture-image2-v2.png",
        "图 4-1  聚变数字孪生参考架构：统一配置、时间、单位、不确定度与证据，把物理、工程、实验和控制连接起来。概念性科学架构图。",
        width=6.35,
    )
    rows = []
    for domain in DOMAIN_ORDER:
        label, en, _ = DOMAIN_META[domain]
        domain_entries = [item for item in entries if domain in entry_domains(item)]
        primary_count = sum(item["primaryDomain"] == domain for item in domain_entries)
        related_count = sum(is_related_work(item, domain) for item in domain_entries)
        evidence = Counter(item["evidenceLevel"] for item in domain_entries)
        deployment = Counter(item["deploymentLevel"] or "未判级" for item in domain_entries)
        rows.append([
            label,
            str(primary_count),
            str(related_count),
            str(len(domain_entries)),
            ", ".join(f"{k or '未标注'}:{v}" for k, v in sorted(evidence.items())),
            ", ".join(f"{k}:{v}" for k, v in sorted(deployment.items())),
            DOMAIN_SYNTHESIS[domain]["gap"],
        ])
    add_simple_table(doc, ["知识域", "主域", "关联", "关系数", "证据分布", "部署分布", "进入数字孪生的主要缺口"], rows, [1100, 560, 560, 650, 1250, 1450, 3790], font_size=7.1)
    table_citation(doc, f"表 4-1  动态统计来自结构化 JSON。唯一工作 {stats['unique_total']} 个；领域关系 {stats['domain_relationships']} 条（主域 {stats['primary_relationships']} + 关联域 {stats['related_relationships']}）。")


def code_summary(item: dict[str, Any]) -> str:
    labels = OrderedDict()
    for code in item.get("code", []):
        labels[code_status_label(code.get("status"))] = True
    return " / ".join(labels.keys()) or "未记录代码资产"


def add_research_entry(doc: Document, item: dict[str, Any], local_index: int, chapter_number: int, chapter_domain: str):
    related = is_related_work(item, chapter_domain)
    title_prefix = "【关联工作】" if related else ""
    doc.add_heading(f"{chapter_number}.2.{local_index}  {title_prefix}{item['title']}", level=3)
    add_body(doc, f"解决问题：{item['problem']}", bold_prefix="解决问题：")
    add_body(doc, f"技术路径：{item['approach']}", bold_prefix="技术路径：")
    metadata = [
        ["projectId", item["projectId"]],
        ["年份 / 机构", f"{item['year']} · {item['organization']}"],
        ["章节关系", (f"关联工作；主域为 {domain_display(item['primaryDomain'])}" if related else f"主域工作；关联域：{'、'.join(domain_display(d) for d in item['relatedDomains']) or '无'}")],
        ["适配 / 验证装置", "；".join(item["devices"]) or "未标注"],
        ["证据等级", f"{item['evidenceLevel'] or '未标注'} · {evidence_label(item['evidenceLevel'])}"],
        ["部署等级", f"{item['deploymentLevel'] or '—'} · {deployment_label(item['deploymentLevel'])}"],
        ["代码关系", code_summary(item)],
    ]
    add_simple_table(doc, ["字段", "内容"], metadata, [1800, 7560], header_fill=DEEP, font_size=8.2)
    table_citation(doc, "条目元数据：以原始论文、官方来源和作者仓库为准。")
    add_body(doc, f"验证与效果：{item['evidence']}", bold_prefix="验证与效果：")
    add_body(doc, f"数据与成熟度：{item['data']}；{item['maturity']}", bold_prefix="数据与成熟度：")
    add_body(doc, f"主要局限：{item['limitations']}", bold_prefix="主要局限：", color="704B32")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("论文 / 原始来源：")
    set_run_font(r, size=9.2, color=DEEP_DARK, bold=True)
    for index, paper in enumerate(item["papers"]):
        if index:
            sep = p.add_run("；")
            set_run_font(sep, size=9.2, color=MUTED)
        paper_text = f"[{paper_source_label(paper.get('sourceType'))}] {paper['title']} ({paper['year']})"
        if paper["url"]:
            add_hyperlink(p, paper_text, paper["url"])
        else:
            paper_run = p.add_run(paper_text)
            set_run_font(paper_run, size=9.2, color=INK)

    for code in item["code"]:
        cp = doc.add_paragraph()
        cp.paragraph_format.left_indent = Inches(0.188)
        cp.paragraph_format.first_line_indent = Inches(-0.188)
        cp.paragraph_format.space_after = Pt(3)
        marker = cp.add_run("• ")
        set_run_font(marker, size=9, color=FUSION, bold=True)
        if code["url"]:
            add_hyperlink(cp, code["name"], code["url"])
        else:
            name = cp.add_run(code["name"])
            set_run_font(name, size=9, color=INK, bold=True)
        qualifiers = " · ".join((
            code_status_label(code.get("status")),
            artifact_type_label(code.get("artifactType")),
            access_label(code.get("access")),
        ))
        status = cp.add_run(f"  [{qualifiers}]  {code['relationship']}")
        set_run_font(status, size=8.8, color=MUTED)


def add_domain_chapter(doc: Document, domain: str, entries: list[dict[str, Any]], figures_dir: Path, chapter_number: int):
    label, en, figure_name = DOMAIN_META[domain]
    add_kicker(doc, f"{chapter_number:02d} / {en}", page_break_before=True)
    doc.add_heading(f"{chapter_number}  {label}", level=1)
    synthesis = DOMAIN_SYNTHESIS[domain]
    add_body(doc, f"需求目标：{synthesis['need']}", bold_prefix="需求目标：")
    add_body(doc, f"当前前沿：{synthesis['state']}", bold_prefix="当前前沿：")
    add_body(doc, f"关键差距：{synthesis['gap']}", bold_prefix="关键差距：", color="704B32")
    add_figure(
        doc,
        figures_dir / figure_name,
        f"图 {chapter_number}-1  {label}的智能原生核心结构。概念性科学架构图，用于说明信息和计算关系，不代表特定装置计算结果。",
        width=6.35,
    )
    domain_entries = [item for item in entries if domain in entry_domains(item)]
    domain_entries.sort(key=lambda item: (is_related_work(item, domain), str(item.get("year", "")), item["title"]))
    if not domain_entries:
        add_callout(doc, "证据空白", "当前核验语料中尚无足够聚变专属公开工作。报告保留该知识域，用于明确技术迁移方向与未来补充入口。", fill=PALE_ORANGE, accent=FUSION)
        return

    index_heading = doc.add_heading(f"{chapter_number}.1  研究分布与快速索引", level=2)
    index_heading.paragraph_format.page_break_before = True
    primary_count = sum(item["primaryDomain"] == domain for item in domain_entries)
    related_count = len(domain_entries) - primary_count
    add_callout(doc, "本章计数口径", f"本章共有 {len(domain_entries)} 条领域关联：{primary_count} 个主域工作、{related_count} 个关联工作。关联工作会在标题中显式标注，并不代表新增唯一项目。", fill=PALE_PURPLE, accent=PURPLE)
    index_rows = []
    for item in domain_entries:
        title_prefix = "【关联工作】" if is_related_work(item, domain) else ""
        index_rows.append([
            f"{title_prefix}{item['title']}",
            "；".join(item["devices"][:2]),
            f"{item['evidenceLevel'] or '—'} / {item['deploymentLevel'] or '—'}",
            code_summary(item),
        ])
    add_simple_table(doc, ["工作", "适配 / 验证装置", "证据 / 部署", "代码"], index_rows, [3200, 2300, 1600, 2260], font_size=7.4)
    table_citation(doc, f"表 {chapter_number}-1  {label}研究索引；E 与 D 是相互独立的两个轴，详细适用边界见后续条目。")
    doc.add_heading(f"{chapter_number}.2  代表工作详解", level=2)
    for idx, item in enumerate(domain_entries, start=1):
        add_research_entry(doc, item, idx, chapter_number, domain)
    if domain == "data":
        add_callout(
            doc,
            "跨域接口提示",
            "TokaMind 在本报告中以数据基座为主域，同时关联物理模拟、智能诊断与人机交互。"
            "网页在四个领域筛选中均可检索到该工作，但唯一工作统计只计一次；其公开证据为装置离线数据与研究服务，"
            "不等同于控制室助手或装置闭环控制。",
            fill=PALE_PURPLE,
            accent=PURPLE,
        )


def add_device_adoption(doc: Document, entries: list[dict[str, Any]], chapter_number: int):
    heading = doc.add_heading(f"{chapter_number}  装置采用与跨装置迁移", level=1)
    heading.paragraph_format.page_break_before = True
    add_body(doc, "聚变 AI 的可信度必须回答‘在哪台装置、什么工况、使用何种信号、是否闭环、失败时如何退化’。同一算法在离线历史数据上表现良好，不意味着能迁移到另一台装置，更不意味着可以直接服务 ITER、SPARC 或 DEMO。")
    rows = [
        ["TCV", "磁位形 RL、轨迹预测与优化", "E4", "高质量仿真器与明确约束是 sim-to-real 前提"],
        ["DIII-D", "破裂/撕裂模、PACMAN、避免控制", "E4", "已具闭环证据；需验证燃烧等离子体与长期运行"],
        ["KSTAR", "不稳定性避免、跨装置 ELM 预测", "E4 / E2", "跨工况稳定性和可解释边界仍需扩大"],
        ["HL-3 / SUNIST-2", "FusionMAE、诊断重建与多任务表征", "E2–E3", "预训练仍偏装置特定，公开数据和权重有限"],
        ["MAST / MAST-U", "开放数据、TokaMark/TokaMind、事件预测", "E2", "开放基准正在形成；闭环证据不足"],
        ["JET / EAST / C-Mod", "跨装置破裂、输运代理、数据管线", "E2–E4", "信号定义、壁条件和运行制度导致域偏移"],
        ["SPARC", "TORAX 场景搜索、热负荷控制研发", "E1", "装置运行前属于仿真与官方研发计划"],
        ["ITER / DEMO", "目标适配、设计优化与安全场景研究", "E0–E1", "缺少同尺度运行数据，必须可证外推"],
    ]
    add_simple_table(doc, ["装置", "代表方向", "最高公开证据", "采用边界"], rows, [1500, 2800, 1500, 3560], font_size=8.2)
    table_citation(doc, f"表 {chapter_number}-1  装置证据摘要；各研究的具体等级以条目为准。")
    doc.add_heading(f"{chapter_number}.1  跨装置迁移的六个检查点", level=2)
    for item in (
        "物理相似性：无量纲参数、壁材料、几何、执行器和诊断是否处于相近域。",
        "语义一致性：同名信号的定义、单位、坐标、采样率、滤波与时间对齐是否一致。",
        "标签一致性：破裂、不稳定性、异常和健康状态的定义是否由同一规则生成。",
        "校准与 UQ：概率是否在目标装置重新校准，训练域外是否有可靠拒绝机制。",
        "实时与接口：最坏时延、丢包、传感器故障和执行器限幅是否进入验证。",
        "安全与责任：AI 失败时独立控制/保护是否保持安全，谁批准模型版本与适用域。",
    ):
        add_bullet(doc, item)


def add_twin_gap(doc: Document, figures_dir: Path, chapter_number: int):
    heading = doc.add_heading(f"{chapter_number}  现有 AI 工作距离数字孪生还缺什么", level=1)
    heading.paragraph_format.page_break_before = True
    add_body(doc, "单个 AI 模型只解决映射问题；数字孪生必须持续代表一个具体资产或装置在特定配置和时间下的状态，并能把模型、数据、实验、工程判断与决策证据闭环起来。当前最主要的差距不是再训练更大的网络，而是把模型嵌入可验证的系统工程。")
    add_figure(
        doc,
        figures_dir / "roadmap-image2-v2.png",
        f"图 {chapter_number}-1  从模型与集成模拟走向运行数字孪生的能力阶梯。概念性路线图。",
        width=6.35,
    )
    gaps = [
        ["权威状态", "许多模型处理历史切片，未维护‘此时此刻、此配置’的唯一状态。", "建立时间同步状态总线、质量标记与状态估计服务。"],
        ["配置与数字线程", "训练数据常缺壁状态、几何版本、校准、控制器和维护记录。", "把配置基线、校准、软件/权重版本关联到每个炮次和仿真。"],
        ["多保真一致性", "代理、高保真代码、工程求解器与实验使用不同变量和误差口径。", "定义公共接口、映射、误差预算和跨模型一致性测试。"],
        ["在线校准", "多数模型离线冻结，无法区分慢漂移与真实物理变化。", "引入影子更新、漂移检测、候选模型和可回退发布。"],
        ["不确定度", "点预测多，校准区间、相关误差和决策风险传播少。", "把 UQ 作为接口字段，并传播到控制、设计和维护决策。"],
        ["失效与降级", "正常数据集评测多，传感器失效、通信丢包和执行器故障少。", "建立故障注入、最坏时延、OOD 与安全降级测试库。"],
        ["实时工程", "平均推理速度不等于可部署；确定性、内存和调度抖动常未测。", "定义 WCET、资源上限、实时平台和 HIL 证据。"],
        ["安全论证", "准确率不能直接转化为安全证据。", "保持独立保护，形成需求—危害—测试—版本—批准证据链。"],
        ["人机协同", "解释多停留在图表，未验证操作者是否理解和正确行动。", "开展人因实验、告警设计、责任分配和程序一致性验证。"],
        ["全厂闭环", "大多数工作局限于等离子体或单子系统。", "构建热、电、氚、维护、经济和安全约束的全厂协调层。"],
    ]
    add_simple_table(doc, ["差距", "当前常见状态", "数字孪生改进"], gaps, [1500, 3700, 4160], font_size=7.9)
    table_citation(doc, f"表 {chapter_number}-1  AI 研究原型到可运行数字孪生的系统性差距。")


def add_fusiondigital_roadmap(doc: Document, chapter_number: int):
    heading = doc.add_heading(f"{chapter_number}  FusionDigital 分阶段总体路线", level=1)
    heading.paragraph_format.page_break_before = True
    add_callout(doc, "路线原则", "先形成可测量价值和可重复证据，再增加跨域耦合与权限。任何阶段都不以‘模型演示成功’替代实验、工程和安全验收。", fill=PALE_ORANGE, accent=FUSION)
    phases = [
        ["P0 证据基座", "0–6 个月", "炮次/仿真数据谱系、统一时间与配置、评测集、模型卡", "关键数据可追溯；基线可重复；权限和发布流程可审计"],
        ["P1 控制服务", "3–12 个月", "DINA / MEQ 服务化、状态估计、快速代理、离线场景比较", "接口稳定；误差/时延有预算；仿真回放优于基线"],
        ["P2 工程代理", "6–18 个月", "商业软件电磁场/力、热管理代理；力、位移、应变、温度验证", "代理误差在适用域内达标；应力反演口径明确；数据闭环可复现"],
        ["P3 影子孪生", "12–30 个月", "物理—控制—工程预测并行运行；UQ、OOD、故障注入", "不影响装置；能提前量化风险；失败可自动降级"],
        ["P4 操作副驾驶", "18–36 个月", "实验检索、场景设计、仿真编排、证据化解释", "回答有引用；工具白名单；无越权写操作；专家节省可量化"],
        ["P5 有限自治", "30–60+ 个月", "狭窄工况内的闭环策略、全厂协调与维护优化", "独立 V&V；HIL/装置证据；安全层独立；权限可撤销"],
    ]
    add_simple_table(doc, ["阶段", "时间", "主要交付", "通过门"], phases, [1450, 1200, 3300, 3410], font_size=8.0)
    table_citation(doc, f"表 {chapter_number}-1  推荐路线；时间以数据、装置窗口和合作资源为条件，不构成刚性承诺。")
    doc.add_heading(f"{chapter_number}.1  建议的五个联合研发工作包", level=2)
    packages = [
        ("WP1 · DINA / MEQ + 学习残差", "把现有控制服务形成标准化输入/输出、版本、回放和不确定度接口；AI 首先学习残差或参数，而不是替换物理核心。"),
        ("WP2 · 电磁力—热管理工程代理", "由高校/商业软件团队生成 DOE 样本与高保真结果；以力、位移、应变、温度实测验证，明确应力为模型反演量。"),
        ("WP3 · 诊断状态表征", "选择少量高价值诊断，建立缺失信号、异常、跨炮次漂移和虚拟诊断基准，逐步形成装置状态向量。"),
        ("WP4 · 仿真与证据智能体", "仅在离线白名单工具中自动建立算例、提交任务、检查收敛、汇总结果与引用，所有写操作保留批准。"),
        ("WP5 · 影子运行与 VVUQ", "建立与实验同步的影子孪生，持续计算预测误差、覆盖率、OOD、时延与故障注入结果，作为权限升级依据。"),
    ]
    for title, text in packages:
        add_callout(doc, title, text, fill=PALE, accent=DEEP)


def add_collaboration(doc: Document, chapter_number: int):
    heading = doc.add_heading(f"{chapter_number}  联合研发与专家协作清单", level=1)
    heading.paragraph_format.page_break_before = True
    add_body(doc, "与高校、装置团队、商业软件供应商或实验专家合作时，最重要的不是先确定网络结构，而是共同冻结问题、变量、工况、证据和交付边界。以下清单可直接用于技术交流。")
    rows = [
        ["物理专家", "适用域、守恒/稳定性约束、父模型可信度、关键无量纲量", "确认哪些外推不可接受"],
        ["实验团队", "传感器定义、校准、时间同步、故障记录、实验窗口", "独立留出炮次与验证方案"],
        ["工程仿真团队", "几何/材料/载荷基线、网格收敛、接触/边界、商业软件版本", "DOE、误差预算与高保真样本"],
        ["控制团队", "周期、时延、限幅、共享执行器、保护与降级", "SIL/HIL/影子测试和接口契约"],
        ["AI 团队", "数据切分、模型、UQ、OOD、解释、版本和监控", "模型卡、复现实验与失效案例"],
        ["V&V / 安全", "危害、需求、独立测试、变更影响、批准与回退", "证据矩阵和权限门"],
        ["平台 / 运维", "数据访问、算力、调度、观测、审计、备份与服务等级", "可运行、可追责、可维护的部署"],
    ]
    add_simple_table(doc, ["角色", "必须共同回答", "最小交付"], rows, [1500, 4800, 3060], font_size=8.1)
    table_citation(doc, f"表 {chapter_number}-1  聚变智能原生跨学科协作的最小责任面。")


def add_appendices(doc: Document, entries: list[dict[str, Any]]):
    heading = doc.add_heading("附录 A  代码库索引", level=1)
    heading.paragraph_format.page_break_before = True
    repos: OrderedDict[str, dict[str, str]] = OrderedDict()
    for item in entries:
        for code in item["code"]:
            key = code["url"] or f"unpublished::{item['projectId']}::{code['name']}"
            if key not in repos:
                repos[key] = {
                    "name": code["name"],
                    "url": code["url"] or "—",
                    "status": code_status_label(code.get("status")),
                    "asset": f"{artifact_type_label(code.get('artifactType'))} / {access_label(code.get('access'))}",
                    "domain": domain_display(item["primaryDomain"]),
                    "relationship": code["relationship"],
                }
    rows = [[repo["name"], repo["domain"], repo["status"], repo["asset"], repo["url"], repo["relationship"]] for repo in repos.values()]
    add_simple_table(doc, ["代码 / 工具", "主域", "关系", "资产 / 访问", "链接", "说明"], rows, [1500, 900, 1150, 1450, 2000, 2360], font_size=6.6)
    table_citation(doc, "代码索引同时以 CSV 和 JSON 交付；链接可能因项目迁移而变化，维护时应复核。")

    doc.add_heading("附录 B  论文与原始来源索引", level=1)
    sources: OrderedDict[str, dict[str, Any]] = OrderedDict()
    for item in entries:
        for paper in item["papers"]:
            source_key = paper["url"] or f"unspecified::{paper['title']}::{paper['year']}"
            sources.setdefault(source_key, paper)
    for index, source in enumerate(sources.values(), start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        number = p.add_run(f"[{index}] ")
        set_run_font(number, size=8.5, color=FUSION, bold=True)
        title = p.add_run(f"[{paper_source_label(source.get('sourceType'))}] {source['title']}. {source['venue']} ({source['year']}). ")
        set_run_font(title, size=8.2, color=INK)
        if source["url"]:
            add_hyperlink(p, source["url"], source["url"])

    doc.add_heading("附录 C  结构化数据维护说明", level=1)
    add_body(doc, "随报告交付的 JSON 是网页和报告的共同数据源。新增条目应保留稳定 projectId，并补齐 primaryDomain、relatedDomains、title、year、organization、problem、approach、devices、evidenceLevel、deploymentLevel、evidence、papers、code、data、maturity、limitations 和 tags。论文宜标 sourceType；代码资产宜标 status、artifactType 与 access。旧字段 id/domain 仍可读取，但不应作为新数据的首选口径。")
    for item in (
        "论文链接优先使用 DOI、期刊、arXiv、OSTI、IAEA 或机构官方页面。",
        "代码必须说明与论文的关系；通用框架只能标记为官方使能工具。",
        "证据等级与部署等级分别维护：只按公开材料直接证明的最高值填写，二者不得自动互推。",
        "每次更新记录检索日期，并重新运行链接、JSON、构建和渲染检查。",
    ):
        add_bullet(doc, item)


def build_report(data_path: Path, figures_dir: Path, output_path: Path):
    payload = json.loads(data_path.read_text(encoding="utf-8"))
    raw_entries = payload.get("entries") or payload.get("projects") or []
    entries = merge_unique_works([normalize_entry(item, index) for index, item in enumerate(raw_entries, start=1)])
    domain_relationships = sum(len(entry_domains(item)) for item in entries)
    stats = {
        "unique_total": len(entries),
        "total": len(entries),  # Backward-compatible alias used by older custom sections.
        "domain_relationships": domain_relationships,
        "primary_relationships": len(entries),
        "related_relationships": domain_relationships - len(entries),
        "device_evidence": sum(item["evidenceLevel"] in {"E2", "E3", "E4"} for item in entries),
        "direct_code": sum(any(code["status"] == "official-direct" for code in item["code"]) for item in entries),
        "code_counts": dict(Counter(code["status"] for item in entries for code in item["code"])),
        "deployment_counts": dict(Counter(item["deploymentLevel"] or "未单独判级" for item in entries)),
    }

    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    configure_section(section)
    configure_headers_footers(section)
    doc.core_properties.title = "聚变智能原生技术图谱"
    doc.core_properties.subject = "九个聚变数字孪生知识域的人工智能研究、论文、代码与实施路线"
    doc.core_properties.author = "新奥聚变人工智能团队 / FusionDigital"
    doc.core_properties.keywords = "FusionDigital, fusion digital twin, AI-native, tokamak, machine learning, AI agents"

    add_cover(doc, figures_dir, stats)
    add_manual_toc(doc)
    add_executive_summary(doc, entries, stats)
    add_primer(doc, figures_dir)
    add_methodology(doc, stats)
    add_landscape_overview(doc, entries, figures_dir, stats)

    # Domain chapters are numbered 5–13 because chapter 4 is the overview.
    for chapter, domain in enumerate(DOMAIN_ORDER, start=5):
        add_domain_chapter(doc, domain, entries, figures_dir, chapter)
    add_device_adoption(doc, entries, 14)
    add_twin_gap(doc, figures_dir, 15)
    add_fusiondigital_roadmap(doc, 16)
    add_collaboration(doc, 17)
    add_appendices(doc, entries)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"saved={output_path}")
    print(f"unique_works={len(entries)} domain_relationships={domain_relationships} paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} inline_shapes={len(doc.inline_shapes)}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--data", type=Path, required=True)
    parser.add_argument("--figures-dir", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    build_report(args.data, args.figures_dir, args.output)


if __name__ == "__main__":
    main()
