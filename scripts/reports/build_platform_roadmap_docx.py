from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ACCENT = "E86F2A"
TEAL = "167C6A"
INK = "18221E"
MUTED = "5E6E67"
LIGHT = "F2F4F3"
LINE = "D7DEDA"
WARN = "FFF4EA"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_borders(table, color=LINE, size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths: list[int], total=9360, indent=120) -> None:
    if sum(widths) != total:
        raise ValueError(f"table widths {widths} do not sum to {total}")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def set_east_asia_font(run, font="Microsoft YaHei") -> None:
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("FusionDigital 技术路线  ·  ")
    set_east_asia_font(run)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def add_rule(paragraph, color=ACCENT, size="18") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_runs(paragraph, text: str) -> None:
    # Lightweight Markdown emphasis/link rendering.
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_east_asia_font(run)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.color.rgb = RGBColor.from_string(TEAL)
        else:
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            run = paragraph.add_run(f"{label}（{url}）")
            run.font.color.rgb = RGBColor.from_string(TEAL)
            run.underline = True
        set_east_asia_font(run, "Consolas" if token.startswith("`") else "Microsoft YaHei")
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_east_asia_font(run)


def add_callout(doc: Document, title: str, body: str, fill=WARN) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [9360], indent=180)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 130, 180, 130, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_east_asia_font(r)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(ACCENT)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_runs(p, body)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_architecture_visual(doc: Document) -> None:
    table = doc.add_table(rows=5, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [9360], indent=160)
    entries = [
        ("公开投影面", "Sites / vinext / Worker / D1 / R2 · 公开知识、三维、EFIT 与带引用问答", "EAF4F1"),
        ("↓  已审查、签名、版本化的数据产品", "", "FFFFFF"),
        ("内网科学平台面", "PostgreSQL / S3 / MDSplus Gateway / CAD·CAE / Simulation / Knowledge / Agent Tool Broker", "EEF3F8"),
        ("↓  经验证且审批的参数包；不是在线控制命令", "", "FFFFFF"),
        ("实验实时面", "DAQ / PCS / MARTe2 / RT Linux / Interlock / Protection", "FFF1E7"),
    ]
    for row, (title, body, fill) in zip(table.rows, entries):
        cell = row.cells[0]
        set_cell_shading(cell, fill)
        set_cell_margins(cell, 105, 160, 105, 160)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        set_east_asia_font(r)
        r.bold = True
        r.font.size = Pt(10 if body else 9)
        r.font.color.rgb = RGBColor.from_string(INK if body else MUTED)
        if body:
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(body)
            set_east_asia_font(r)
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor.from_string(MUTED)
    set_table_borders(table, "FFFFFF", "2")
    doc.add_paragraph()


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    base, remainder = divmod(9360, cols)
    widths = [base + (1 if index < remainder else 0) for index in range(cols)]
    set_table_geometry(table, widths)
    set_table_borders(table)
    for i, values in enumerate(rows):
        prevent_row_split(table.rows[i])
        if i == 0:
            set_repeat_table_header(table.rows[i])
        for j in range(cols):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_shading(cell, LIGHT if i == 0 else "FFFFFF")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            add_runs(p, values[j] if j < len(values) else "")
            for run in p.runs:
                run.font.size = Pt(8.2)
                if i == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Title", 31, INK, 0, 10),
        ("Subtitle", 13, MUTED, 0, 14),
        ("Heading 1", 20, INK, 20, 8),
        ("Heading 2", 14, TEAL, 14, 6),
        ("Heading 3", 11, ACCENT, 10, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("FUSIONDIGITAL · PLATFORM BLUEPRINT")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(ACCENT)
    set_east_asia_font(r)
    add_rule(p)

    doc.add_paragraph().paragraph_format.space_after = Pt(24)
    p = doc.add_paragraph(style="Title")
    p.add_run("FusionDigital\n整体技术路线图")
    for run in p.runs:
        set_east_asia_font(run)
    p = doc.add_paragraph(style="Subtitle")
    p.add_run("从公开展示原型到可复现的科学、工程与控制协作平台")
    for run in p.runs:
        set_east_asia_font(run)

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(meta, [1512, 7848], indent=110)
    labels = [
        ("版本", "2026-08-15"),
        ("状态", "建议架构 / 实施路线"),
        ("范围", "知识、三维、LLM/智能体、物理与工程仿真、诊断、控制、数据基座"),
        ("适用", "FusionDigital 下一阶段建设与跨团队技术沟通"),
    ]
    for row, (label, value) in zip(meta.rows, labels):
        prevent_row_split(row)
        for cell in row.cells:
            set_cell_margins(cell, 90, 110, 90, 110)
        set_cell_shading(row.cells[0], LIGHT)
        row.cells[0].width = Inches(1.05)
        row.cells[1].width = Inches(5.45)
        p0 = row.cells[0].paragraphs[0]
        p1 = row.cells[1].paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        p1.paragraph_format.space_after = Pt(0)
        r = p0.add_run(label)
        set_east_asia_font(r)
        r.bold = True
        r.font.size = Pt(9)
        add_runs(p1, value)
        for r in p1.runs:
            r.font.size = Pt(9)
    set_table_borders(meta)

    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    add_callout(
        doc,
        "核心决策",
        "保留现有站点为公开投影面；新建内网科学平台承载 MDSplus、NAS、对象存储、CAD/CAE、仿真和诊断；PCS、联锁与保护留在隔离的实时域。三者只通过版本化、可审查、可重放的合同连接。",
    )
    doc.add_page_break()


def parse_markdown_into_doc(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    i = 0
    first_heading = True
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            first_heading = False
            i += 1
            continue
        if line.startswith("**版本：") or line.startswith("**范围："):
            i += 1
            continue
        first_heading = False

        if line.startswith("## "):
            title = re.sub(r"^##\s+", "", line)
            p = doc.add_paragraph(title, style="Heading 1")
            if title.startswith("3. 目标架构"):
                add_architecture_visual(doc)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_paragraph(re.sub(r"^###\s+", "", line), style="Heading 2")
            i += 1
            continue
        if line.startswith("#### "):
            doc.add_paragraph(re.sub(r"^####\s+", "", line), style="Heading 3")
            i += 1
            continue

        if line.startswith("```"):
            lang = line[3:].strip()
            block: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1
            if lang == "mermaid":
                continue
            table = doc.add_table(rows=1, cols=1)
            set_table_geometry(table, [9360], indent=150)
            cell = table.cell(0, 0)
            set_cell_shading(cell, "F6F7F7")
            set_cell_margins(cell, 120, 150, 120, 150)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run("\n".join(block))
            r.font.name = "Consolas"
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor.from_string(INK)
            continue

        if line.startswith("|"):
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                values = [item.strip() for item in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", value) for value in values):
                    rows.append(values)
                i += 1
            add_markdown_table(doc, rows)
            continue

        if line.startswith("> "):
            add_callout(doc, "当前系统定位", line[2:].strip(), "EEF3F8")
            i += 1
            continue

        if re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            add_runs(p, line)
            i += 1
            continue

        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            add_runs(p, line[2:].strip())
            i += 1
            continue

        paragraph_lines = [line]
        i += 1
        while i < len(lines):
            candidate = lines[i].strip()
            if not candidate or candidate.startswith(("#", "|", "```", "> ", "- ")) or re.match(r"^\d+\.\s+", candidate):
                break
            paragraph_lines.append(candidate)
            i += 1
        p = doc.add_paragraph()
        add_runs(p, " ".join(paragraph_lines))


def build(source: Path, output: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)
    section.orientation = WD_ORIENT.PORTRAIT
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.text = "FUSIONDIGITAL  /  PLATFORM ARCHITECTURE"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        set_east_asia_font(run)
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(TEAL)
    add_rule(header, TEAL, "6")
    add_page_number(section.footer.paragraphs[0])

    add_cover(doc)
    parse_markdown_into_doc(doc, source.read_text(encoding="utf-8"))

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "FusionDigital 整体技术路线图"
    doc.core_properties.subject = "知识、三维、仿真、诊断、控制、数据与智能体平台架构"
    doc.core_properties.author = "FusionDigital"
    doc.core_properties.keywords = "FusionDigital, MDSplus, CAD, CAE, EFIT, DINA, MEQ, digital twin"
    doc.save(output)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("usage: build_platform_roadmap_docx.py SOURCE.md OUTPUT.docx")
    build(Path(sys.argv[1]).resolve(), Path(sys.argv[2]).resolve())
