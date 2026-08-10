from __future__ import annotations

import argparse
import json
import re
from collections import Counter
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from build_ai_native_report import (
    CYAN,
    DEEP,
    DEEP_DARK,
    FUSION,
    INK,
    LINE,
    MUTED,
    PALE,
    PALE_ORANGE,
    PALE_PURPLE,
    PURPLE,
    WHITE,
    add_body,
    add_bullet,
    add_callout,
    add_figure,
    add_hyperlink,
    add_kicker,
    add_metric_strip,
    add_page_number,
    add_simple_table,
    configure_section,
    configure_styles,
    set_run_font,
)


TASK_ORDER = [f"T{index}" for index in range(10)]
TASK_INTRO_FALLBACK = {
    "T0": "状态估计不是普通离线诊断，而是所有闭环共享的实时服务。它必须在规定时限内发布带时间戳、质量标志、模型版本和不确定度的控制状态，并能在诊断缺失、漂移和模型分歧时显式降级。",
    "T1": "启动、电流和磁通控制覆盖击穿、烧穿、升流、平顶、非感应维持与受控降流。它连接电源、磁体、真空室涡流、等离子体与场景轨迹，是 DINA/MEQ 数字孪生最自然的第一条主链。",
    "T2": "位置、位形与边界控制包含快速垂直稳定、形状、间隙、X 点、打击点和先进偏滤器拓扑。快慢线圈分工、响应矩阵秩、饱和与线圈力约束比单一几何误差更重要。",
    "T3": "剖面与场景控制面向 q/电流、温度、密度、压力、旋转、杂质和非感应电流分数。核心矛盾是输运记忆、诊断稀疏、多执行器副作用与实时模型保真度。",
    "T4": "稳定性控制不是一个算法，而是 NTM、RWM、锯齿、AE、ELM、误差场和约束模式转换等多类检测—定位—执行链。它们可共享事件总线和监督器，但应保持机制与失效模式的独立证据。",
    "T5": "排热控制必须同时看核心性能、辐射、脱靶、热流、材料温度、杂质和壁库存。等离子体可控不等于部件寿命可接受，控制孪生需要连接边缘物理模型、红外/光谱实测和工程限值。",
    "T6": "性能、功率与燃烧控制从 β、储能和中子率逐步走向聚变功率、Q、alpha 加热、燃料比和氦灰。电厂层还需以更慢时间尺度协调热循环、厂用电与电网，但不能把全部职责塞入毫秒级 PCS。",
    "T7": "破裂相关控制应区分预测、避免、恢复、受控终止和缓解触发。数字孪生可以为触发边界和策略形成证据，但不能替代独立保护链；E4 装置闭环也不自动构成 D5 安全关键批准。",
    "T8": "集成控制显式处理目标冲突、共享执行器、约束、优先级、重构和降级。近期最有价值的数字孪生能力是候选动作比较、执行器能力模型和影子验证，而不是一步到位的全局自治。",
    "T9": "PCS 承载实时算法，也负责配置、脉冲状态机、时钟、数据传输、执行器接口、日志、回放、权限和软件质量。其验证需要从单元测试逐级走向 SIL/HIL、影子、装置闭环和持续运行。",
}

NOTE_TASK_HEADING = {
    "T0": "T0：状态估计与实时诊断",
    "T1": "T1 与 T2：磁控制主干",
    "T2": "T1 与 T2：磁控制主干",
    "T3": "T3：剖面与场景控制",
    "T4": "T4：不稳定性与约束模式控制",
    "T5": "T5：热负荷、辐射、粒子与壁控制",
    "T6": "T6：性能、功率与燃烧控制",
    "T7": "T7：失稳避免、安全终止与保护边界",
    "T8": "T8：多执行器协调与集成控制",
    "T9": "T9：PCS 与验证基础设施",
}


def add_control_kicker(doc: Document, text: str, *, page_break_before: bool = False):
    """Render a section kicker that cannot be orphaned from its heading."""
    paragraph = add_kicker(doc, text, page_break_before=page_break_before)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def configure_control_header(section) -> None:
    first = section.first_page_header
    p = first.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p.add_run("Fusion")
    set_run_font(r1, size=9, color=FUSION, bold=True)
    r2 = p.add_run("Digital")
    set_run_font(r2, size=9, color=CYAN, bold=True)

    header = section.header
    hp = header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run("FusionDigital  /  聚变集成控制与 PCS 技术图谱")
    set_run_font(run, size=8, color=MUTED, bold=True)
    p_pr = hp._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), LINE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    label = fp.add_run("新奥聚变人工智能团队  ·  ")
    set_run_font(label, size=8, color=MUTED)
    add_page_number(fp)


def add_cover(doc: Document, figures_dir: Path, stats: dict[str, Any], as_of: str) -> None:
    add_control_kicker(doc, "FUSIONDIGITAL / INTEGRATED CONTROL RESEARCH ATLAS / 2026")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("聚变集成控制与 PCS")
    set_run_font(r, size=28, color=DEEP_DARK, bold=True)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(8)
    r2 = p2.add_run("技术图谱研究报告")
    set_run_font(r2, size=28, color=FUSION, bold=True)
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(8)
    r3 = p3.add_run("Integrated Tokamak Control, Plasma Control Systems and Digital-Twin Roadmap")
    set_run_font(r3, size=12, color=DEEP, bold=True)

    add_callout(
        doc,
        "双索引",
        "以 T0–T9 控制任务回答‘控制什么、靠什么测量、如何执行、证据如何’，再以装置与 PCS 回答‘在哪里做过、运行于什么系统、论文和代码在哪里’。",
        fill=PALE_ORANGE,
        accent=FUSION,
    )
    cover_figure = figures_dir / "control-closed-loop-architecture-nature.png"
    if cover_figure.exists():
        add_figure(doc, cover_figure, "图 0-1  聚变集成控制的闭环信息与责任架构：装置、诊断、状态估计、任务控制、执行器分配、PCS 与独立保护。", width=5.0)
    add_metric_strip(doc, [
        (str(stats["total"]), "项关键控制工作"),
        (str(stats["uniquePapers"]), "篇/项原始来源"),
        (str(stats["devices"]), "个装置与 PCS 档案"),
        ("≥50k", "中文正文门槛"),
    ])
    p4 = doc.add_paragraph()
    p4.paragraph_format.space_before = Pt(4)
    p4.paragraph_format.space_after = Pt(0)
    r4 = p4.add_run(f"新奥聚变人工智能团队 · FusionDigital  |  研究截止：{as_of}  |  版本：1.0  |  联系：tianshao1992@gmail.com")
    set_run_font(r4, size=7.6, color=DEEP_DARK, bold=True)


def add_toc(doc: Document, task_meta: dict[str, Any]) -> None:
    add_control_kicker(doc, "CONTENTS / READING MAP", page_break_before=True)
    doc.add_heading("目录与阅读路径", level=1)
    rows = [
        ["01", "执行摘要：从单回路控制走向可证的多速率协同"],
        ["02", "分类、架构、时间尺度与研究方法"],
        ["03", "证据 E0–E4、部署 D1–D5 与代码关系口径"],
        ["04", "集成控制与数字孪生的边界"],
    ]
    for index, task in enumerate(TASK_ORDER, 5):
        rows.append([f"{index:02d}", f"{task}  {task_meta[task]['label']}"])
    rows.extend([
        ["15", "装置与 PCS 双索引"],
        ["16", "PCS、软件栈、接口与验证基础设施"],
        ["17", "从集成模拟/控制到数字孪生还缺什么"],
        ["18", "FusionDigital 分阶段实施路线"],
        ["19", "专题综合调研笔记与完整索引"],
    ])
    add_simple_table(doc, ["章节", "主题"], rows, [1000, 8360], font_size=9.1)
    reading = doc.add_paragraph()
    reading.paragraph_format.space_before = Pt(5)
    reading.paragraph_format.space_after = Pt(0)
    run = reading.add_run("阅读建议：先看第 1–4、15、17、18 章；讨论具体任务时，再进入对应 T0–T9 章节及条目的论文、代码与验证证据。")
    set_run_font(run, size=7.6, color=MUTED)
    doc.add_page_break()


def parse_markdown_sections(path: Path) -> dict[str, str]:
    if not path.exists():
        return {}
    result: dict[str, list[str]] = {}
    heading = "导言"
    result[heading] = []
    for line in path.read_text(encoding="utf-8").splitlines():
        if line.startswith("# "):
            continue
        if line.startswith("## "):
            heading = re.sub(r"^\d+\.\s*", "", line[3:].strip())
            result.setdefault(heading, [])
        else:
            result[heading].append(line)
    return {key: "\n".join(lines).strip() for key, lines in result.items() if "\n".join(lines).strip()}


def normalize_body_text(value: str) -> str:
    """Return the canonical form used for exact paragraph de-duplication."""
    return re.sub(r"\s+", " ", value).strip()


def iter_body_paragraph_text(doc: Document):
    """Yield paragraph text from the document body, including table cells.

    Headers and footers are deliberately excluded: the report-length gate applies
    to the substantive body, not repeated running furniture.
    """
    for paragraph in doc.element.body.iter(qn("w:p")):
        value = "".join(node.text or "" for node in paragraph.iter(qn("w:t")))
        normalized = normalize_body_text(value)
        if normalized:
            yield normalized


def body_text_statistics(doc: Document) -> dict[str, int]:
    """Count raw and exact-paragraph-de-duplicated body text."""
    paragraphs = list(iter_body_paragraph_text(doc))
    unique_paragraphs = list(dict.fromkeys(paragraphs))

    def cjk_count(values: list[str]) -> int:
        return sum(len(re.findall(r"[\u3400-\u4dbf\u4e00-\u9fff]", value)) for value in values)

    return {
        "paragraphs": len(paragraphs),
        "uniqueParagraphs": len(unique_paragraphs),
        "duplicateParagraphs": len(paragraphs) - len(unique_paragraphs),
        "rawCjk": cjk_count(paragraphs),
        "deduplicatedCjk": cjk_count(unique_paragraphs),
        "rawChars": sum(len(value) for value in paragraphs),
        "deduplicatedChars": sum(len(value) for value in unique_paragraphs),
    }


def add_markdown_text(doc: Document, content: str, heading_level: int = 2) -> None:
    pending: list[str] = []

    def flush() -> None:
        if pending:
            add_body_with_links(doc, " ".join(part.strip() for part in pending if part.strip()))
            pending.clear()

    for raw in content.splitlines():
        line = raw.strip()
        if not line:
            flush()
        elif line.startswith("### "):
            flush()
            doc.add_heading(line[4:].strip(), level=min(heading_level + 1, 3))
        elif line.startswith("## "):
            flush()
            doc.add_heading(re.sub(r"^\d+\.\s*", "", line[3:].strip()), level=heading_level)
        elif line.startswith("# "):
            flush()
        elif re.match(r"^[-*]\s+", line):
            flush()
            add_bullet_with_links(doc, re.sub(r"^[-*]\s+", "", line))
        else:
            pending.append(line)
    flush()


def add_markdown_text_deduplicated(
    doc: Document,
    content: str,
    seen_paragraphs: set[str],
    heading_level: int = 2,
) -> tuple[int, int]:
    """Render Markdown while skipping exact duplicate body paragraphs.

    Headings are retained because they carry document structure. Prose and list
    items are compared after whitespace normalization against both the report
    body already rendered and earlier extended-note files.
    """
    pending: list[str] = []
    added = 0
    skipped = 0

    def add_unique(value: str, *, bullet: bool = False) -> None:
        nonlocal added, skipped
        normalized = normalize_body_text(value)
        if not normalized:
            return
        if normalized in seen_paragraphs:
            skipped += 1
            return
        if bullet:
            add_bullet_with_links(doc, value)
        else:
            add_body_with_links(doc, value)
        seen_paragraphs.add(normalized)
        added += 1

    def flush() -> None:
        if pending:
            add_unique(" ".join(part.strip() for part in pending if part.strip()))
            pending.clear()

    for raw in content.splitlines():
        line = raw.strip()
        if not line:
            flush()
        elif line.startswith("### "):
            flush()
            doc.add_heading(line[4:].strip(), level=min(heading_level + 1, 3))
        elif line.startswith("## "):
            flush()
            doc.add_heading(re.sub(r"^\d+\.\s*", "", line[3:].strip()), level=heading_level)
        elif line.startswith("# "):
            flush()
        elif re.match(r"^[-*]\s+", line):
            flush()
            add_unique(re.sub(r"^[-*]\s+", "", line), bullet=True)
        else:
            pending.append(line)
    flush()
    return added, skipped


def add_body_with_links(doc: Document, value: str, *, size: float = 10.5, color: str = INK):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    cursor = 0
    for match in re.finditer(r"<(?P<url>https?://[^>]+)>|(?P<plain>https?://[^\s，。；）]+)", value):
        if match.start() > cursor:
            run = paragraph.add_run(value[cursor:match.start()])
            set_run_font(run, size=size, color=color)
        url = match.group("url") or match.group("plain")
        add_hyperlink(paragraph, url, url)
        cursor = match.end()
    if cursor < len(value):
        run = paragraph.add_run(value[cursor:])
        set_run_font(run, size=size, color=color)
    return paragraph


def add_bullet_with_links(doc: Document, value: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    match = re.search(r"<(?P<url>https?://[^>]+)>|(?P<plain>https?://[^\s，。；）]+)", value)
    if not match:
        run = p.add_run(value)
        set_run_font(run, size=10, color=INK)
        return
    before = p.add_run(value[:match.start()])
    set_run_font(before, size=10, color=INK)
    url = match.group("url") or match.group("plain")
    add_hyperlink(p, url, url)
    after = p.add_run(value[match.end():])
    set_run_font(after, size=10, color=INK)


def add_executive(doc: Document, entries: list[dict[str, Any]], devices: list[dict[str, Any]], task_meta: dict[str, Any], figures_dir: Path) -> None:
    add_control_kicker(doc, "01 / EXECUTIVE SUMMARY", page_break_before=False)
    doc.add_heading("执行摘要：控制数字孪生的目标是把动作变成可验证的决策", level=1)
    add_body(doc, "聚变装置控制的难点不是缺少单个反馈器，而是目标、模型、诊断、执行器、时延、故障和责任边界同时耦合。位形、剖面、MHD、排热、功率和破裂规避不能各自假设执行器无限可用；PCS、机器保护、中央控制和核安全也不能被一个‘全局智能控制器’替代。")
    add_body(doc, "本报告把用户提出的位形、剖面、不稳定性、热负荷、功率、控制集成和 PCS 重新组织为 T0–T9 十类。T0 状态估计与 T9 PCS/V&V 是横切能力，T1–T8 是被控任务。每项工作同时记录问题、方法、控制架构、时间尺度、传感器、执行器、适配装置、验证、结果、论文、代码和数字孪生意义。")
    add_metric_strip(doc, [
        (str(len(entries)), "项结构化工作"),
        (str(len(devices)), "个装置/PCS 档案"),
        (str(sum(item["evidenceLevel"] == "E4" for item in entries)), "项装置闭环证据"),
        (str(sum(any(code["status"] == "official-direct" for code in item["code"]) for item in entries)), "项官方直接代码"),
    ])
    add_callout(doc, "核心判断", "集成控制距离数字孪生的主要差距，不是再增加一个求解器，而是缺少配置权威、实时状态质量、跨任务执行器契约、系统化 SIL/HIL、持续 VVUQ、失败证据、独立保护边界和全生命周期治理。", fill=PALE_PURPLE, accent=PURPLE)
    time_figure = figures_dir / "control-task-timescale-nature.png"
    if time_figure.exists():
        add_figure(doc, time_figure, "图 1-1  T0–T9 控制任务的典型时间尺度与模型保真度分层。时间范围是工程设计参考，不代表所有装置采用相同周期。")

    doc.add_heading("十类任务的相互关系", level=2)
    rows = [[task, task_meta[task]["label"], task_meta[task]["role"], str(sum(item["primaryTask"] == task for item in entries)), str(sum(task in [item["primaryTask"], *item["relatedTasks"]] for item in entries))] for task in TASK_ORDER]
    add_simple_table(doc, ["ID", "任务", "性质", "主任务工作", "含关联工作"], rows, [650, 3520, 1650, 1770, 1770], font_size=7.8)


def add_methodology(doc: Document, landscape: dict[str, Any], figures_dir: Path) -> None:
    add_control_kicker(doc, "02 / TAXONOMY, ARCHITECTURE & METHOD", page_break_before=False)
    doc.add_heading("分类、架构、时间尺度与研究方法", level=1)
    add_body(doc, "传统分类经常把‘被控物理对象’、‘控制功能’和‘承载软件’放在同一级，例如把位形控制与 PCS 并列。本文采用双轴：纵轴是 T0–T9 任务，横轴是装置、PCS、传感器、执行器、模型与验证层。这样同一篇工作可以有唯一主任务，并明确关联任务，而 PCS 不再被误解成另一个物理回路。")
    add_body(doc, "研究收录优先使用同行评议原始论文、IAEA/机构正式材料、装置官方页面和官方代码仓库。综述只用于发现线索；性能数值回到原始来源。‘使用装置数据’、‘在实时计算机运行’、‘直接影响执行器’分别对应不同证据等级。代码则严格区分论文直接实现、官方使能框架、商业软件与未公开资产。")
    architecture = figures_dir / "control-closed-loop-architecture-nature.png"
    if architecture.exists():
        add_figure(doc, architecture, "图 2-1  分层闭环与责任边界。PCS 使用状态和约束完成受控动作；机器保护与安全系统保持独立权威，同时共享必要状态和事件。")
    doc.add_heading("研究范围与非目标", level=2)
    for item in (
        "范围覆盖托卡马克等离子体与装置级控制、PCS 架构、实时软件、控制设计模型、验证设施和向数字孪生的演进。",
        "不把‘理论可控’写成‘装置已验证’，不把‘接入实时框架’写成‘真实闭环’，不把‘装置闭环’写成‘安全关键资格鉴定’。",
        "开源代码只在与论文或项目关系明确时标记为直接实现；通用求解器、框架和社区复现不得冒充论文原代码。",
        "时间尺度、指标和装置适配均以来源语境为边界；不同硬件与物理尺度下不可机械迁移。",
    ):
        add_bullet(doc, item)


def add_evidence_method(doc: Document, landscape: dict[str, Any], figures_dir: Path) -> None:
    add_control_kicker(doc, "03 / EVIDENCE, DEPLOYMENT & CODE", page_break_before=False)
    doc.add_heading("证据、部署与代码关系：三条独立坐标轴", level=1)
    add_body(doc, "证据等级描述公开材料实际证明了什么；部署等级描述功能如何进入装置工作流；代码关系描述复现资产与论文的距离。三者不可互相自动推导。例如，装置闭环实验属于 E4，但若只是少量专项放电，仍可能只是 D4，而不是经治理批准的 D5 保护功能。")
    rows = [[key, value] for key, value in landscape["evidenceMeta"].items()]
    add_simple_table(doc, ["等级", "证据定义"], rows, [950, 8410], font_size=8.7)
    rows = [[key, value] for key, value in landscape["deploymentMeta"].items()]
    add_simple_table(doc, ["等级", "部署定义"], rows, [950, 8410], font_size=8.7)
    verification = figures_dir / "control-verification-ladder-nature.png"
    if verification.exists():
        add_figure(doc, verification, "图 3-1  控制验证阶梯：从数值基准、历史回放、SIL/HIL 和影子模式，逐级进入受限装置实验、目标工况和持续运行。每一级都需要保存配置与失败证据。")
    doc.add_heading("最低验证信息", level=2)
    for item in (
        "E1：植物模型版本、场景、扰动、传感器/执行器、时延、步长、约束和对照基线。",
        "E2：按放电/实验期的数据切分、数据泄漏检查、缺测/漂移、独立重建与不确定度。",
        "E3：最坏执行时间、抖动、缺帧、超时、数值异常、资源竞争、回退与 HIL 配置。",
        "E4：放电数量、成功率、对照组、操作员/保护系统作用、失败案例和适用工况。",
        "D5：治理批准、责任人、版本冻结、变更控制、持续监视、统计可用率和独立保护边界。",
    ):
        add_bullet(doc, item)


def add_integrated_twin_boundary(doc: Document, synthesis: dict[str, str]) -> None:
    add_control_kicker(doc, "04 / CONTROL VS DIGITAL TWIN", page_break_before=True)
    doc.add_heading("集成控制与数字孪生的边界", level=1)
    wanted = ["集成控制究竟“集成”什么", "数字孪生在控制中的角色", "控制时间尺度与模型保真度"]
    for heading in wanted:
        if heading in synthesis:
            doc.add_heading(heading, level=2)
            add_markdown_text(doc, synthesis[heading], heading_level=3)
    add_callout(doc, "判断标准", "若一个系统只有控制模型和实时曲线，却不能回答模型/配置版本、状态质量、适用域、验证证据、失败回退和动作责任，就仍是先进控制或集成模拟，而不是可信的控制数字孪生。", fill=PALE_ORANGE, accent=FUSION)


def add_source_links(doc: Document, papers: list[dict[str, Any]], codes: list[dict[str, Any]]) -> None:
    doc.add_heading("论文、代码与复现关系", level=3)
    for paper in papers:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        prefix = p.add_run(f"论文｜{paper.get('year') or '—'} · {paper.get('venue') or '原始来源'} · ")
        set_run_font(prefix, size=8.8, color=MUTED)
        add_hyperlink(p, paper["title"], paper["url"])
        if paper.get("doi"):
            suffix = p.add_run(f"  DOI: {paper['doi']}")
            set_run_font(suffix, size=8.2, color=MUTED)
    for code in codes:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        prefix = p.add_run("软件｜")
        set_run_font(prefix, size=8.8, color=MUTED, bold=True)
        if code.get("url"):
            add_hyperlink(p, code["name"], code["url"])
        else:
            name = p.add_run(code["name"])
            set_run_font(name, size=9.2, color=DEEP_DARK, bold=True)
        relation = p.add_run(f"  [{code['status']}; {code.get('access', '未标注')}]  {code['relationship']}")
        set_run_font(relation, size=8.7, color=MUTED)


def add_research_entry(doc: Document, item: dict[str, Any], local_index: int, task: str) -> None:
    title = f"{task}.{local_index:02d}  {item['titleZh']}"
    heading = doc.add_heading(title, level=2)
    heading.paragraph_format.keep_with_next = True
    if item.get("titleEn"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(item["titleEn"])
        set_run_font(run, size=9, color=MUTED, italic=True)
    status = doc.add_paragraph()
    status.paragraph_format.space_before = Pt(0)
    status.paragraph_format.space_after = Pt(5)
    status.paragraph_format.keep_with_next = True
    status_pr = status._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), PALE)
    status_pr.append(shading)
    badge = status.add_run(f"{item['evidenceLevel']} / {item['deploymentLevel']}  ")
    set_run_font(badge, size=8.4, color=DEEP, bold=True)
    metadata = status.add_run(f"{item['year']} · {item['organization']} · 主任务 {item['primaryTask']} · 关联 {', '.join(item['relatedTasks']) or '—'} · 条目 {item['id']}")
    set_run_font(metadata, size=8.4, color=INK)
    first_body = add_body(doc, f"解决问题：{item['problem']}", bold_prefix="解决问题：")
    first_body.paragraph_format.keep_together = True
    add_body(doc, f"方法与控制架构：{item['method']} {item['controlArchitecture']}", bold_prefix="方法与控制架构：")
    add_body(doc, f"时间尺度：{item['timescale']}", bold_prefix="时间尺度：")
    add_body(doc, f"传感器/状态：{'；'.join(item['sensors']) if item['sensors'] else '公开来源未完整列出。'}", bold_prefix="传感器/状态：")
    add_body(doc, f"执行器：{'；'.join(item['actuators']) if item['actuators'] else '公开来源未完整列出。'}", bold_prefix="执行器：")
    add_body(doc, f"适配装置：{'；'.join(item['devices'])}", bold_prefix="适配装置：")
    add_body(doc, f"验证与关键结果：{item['validation']} {item['results']}", bold_prefix="验证与关键结果：")
    add_body(doc, f"成熟度与局限：{item['maturity']} {item['limitations']}", bold_prefix="成熟度与局限：")
    add_body(doc, f"数字孪生意义：{item['twinRelevance']}", bold_prefix="数字孪生意义：")
    add_source_links(doc, item["papers"], item["code"])


def add_task_chapter(
    doc: Document,
    task: str,
    primary_items: list[dict[str, Any]],
    related_items: list[dict[str, Any]],
    task_meta: dict[str, Any],
    synthesis: dict[str, str],
    figures_dir: Path,
    chapter: int,
    include_synthesis: bool = True,
) -> None:
    # Task chapters form one continuous reference section.  Let Word place the
    # next chapter in remaining space when possible; the heading styles keep
    # the kicker/title together and prevent sparse source-only tail pages.
    add_control_kicker(doc, f"{chapter:02d} / {task} / {task_meta[task]['en']}", page_break_before=False)
    doc.add_heading(f"{task}  {task_meta[task]['label']}", level=1)
    add_body(doc, TASK_INTRO_FALLBACK[task])
    note_heading = NOTE_TASK_HEADING[task]
    if include_synthesis and note_heading in synthesis:
        add_markdown_text(doc, synthesis[note_heading], heading_level=2)
    elif note_heading in synthesis:
        add_callout(
            doc,
            "共享综合说明",
            f"“{note_heading}”已在首次使用该主题的前一任务章集中展开；本章不重复相同综合正文，只保留本任务的主工作全文和关联工作速查。",
            fill=PALE,
            accent=DEEP,
        )
    if task == "T9":
        figure = figures_dir / "control-pcs-layers-nature.png"
        if figure.exists():
            add_figure(doc, figure, "图 14-1  PCS 分层架构：从装置 I/O、实时状态和任务控制到监督编排、操作界面与独立保护接口。")
    add_callout(doc, "本章证据", f"全文展开 {len(primary_items)} 项主任务工作，并以速查表列出 {len(related_items)} 项跨任务关联工作。关联条目的完整方法、论文与代码见其主任务章节。", fill=PALE_ORANGE if task in {"T1", "T2", "T4", "T5"} else PALE, accent=FUSION if task in {"T1", "T2", "T4", "T5"} else DEEP)
    doc.add_heading("跨任务关联工作速查", level=2)
    rows = [[item["id"], item["primaryTask"], item["titleZh"], "；".join(item["devices"][:2]), f"{item['evidenceLevel']}/{item['deploymentLevel']}"] for item in related_items]
    if not rows:
        rows = [["—", "—", "当前数据集中没有显式关联工作", "—", "—"]]
    add_simple_table(doc, ["ID", "主任务", "关联工作", "主要装置", "证据/部署"], rows, [1050, 800, 3670, 2600, 1240], font_size=6.9)
    add_body(doc, "说明：本表用于回答该任务还依赖哪些跨域工作；为避免重复扩充篇幅，条目正文及其论文、代码链接只在唯一主任务章节完整展示。")
    for index, item in enumerate(primary_items, 1):
        add_research_entry(doc, item, index, task)


def add_device_chapter(doc: Document, devices: list[dict[str, Any]], entries: list[dict[str, Any]], synthesis: dict[str, str]) -> None:
    add_control_kicker(doc, "15 / DEVICE & PCS INDEX", page_break_before=True)
    doc.add_heading("从装置角度梳理主要控制论文、PCS 与代码", level=1)
    if "从装置角度阅读控制论文" in synthesis:
        add_markdown_text(doc, synthesis["从装置角度阅读控制论文"], heading_level=2)
    add_body(doc, "装置视图不以论文数量排序，而是说明控制栈、主要任务、最高证据和迁移边界。装置名称出现在工作条目中并不自动表示该装置已部署相同控制器；本章仅按档案中的原始来源陈述。")
    for index, device in enumerate(devices, 1):
        heading = doc.add_heading(f"15.{index:02d}  {device['name']}（{device['country']}）", level=2)
        heading.paragraph_format.keep_with_next = True
        status = doc.add_paragraph()
        status.paragraph_format.space_before = Pt(0)
        status.paragraph_format.space_after = Pt(5)
        status.paragraph_format.keep_with_next = True
        status_pr = status._p.get_or_add_pPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), PALE)
        status_pr.append(shading)
        label = status.add_run("装置档案  ")
        set_run_font(label, size=8.4, color=DEEP, bold=True)
        detail = status.add_run(f"{device['organization']} · {device['status']} · 任务 {', '.join(device['primaryTasks']) or '公开资料未完整分类'}")
        set_run_font(detail, size=8.4, color=INK)
        first_body = add_body(doc, f"PCS/控制架构：{device['pcsArchitecture']}", bold_prefix="PCS/控制架构：")
        first_body.paragraph_format.keep_together = True
        add_body(doc, f"时序与周期：{device['timing']}", bold_prefix="时序与周期：")
        add_body(doc, f"主要传感器：{'；'.join(device['sensors']) or '未完整公开。'}", bold_prefix="主要传感器：")
        add_body(doc, f"主要执行器：{'；'.join(device['actuators']) or '未完整公开。'}", bold_prefix="主要执行器：")
        if device["representativeWorks"]:
            add_body(doc, f"代表工作：{'；'.join(device['representativeWorks'])}", bold_prefix="代表工作：")
        add_body(doc, f"成熟度与缺口：{device['maturity']} {device['gaps']}", bold_prefix="成熟度与缺口：")
        add_source_links(doc, device["papers"], device["code"])


def add_pcs_toolchain(doc: Document, figures_dir: Path) -> None:
    add_control_kicker(doc, "16 / PCS & CONTROL TOOLCHAIN", page_break_before=True)
    doc.add_heading("PCS、控制模型、数据接口与验证工具链", level=1)
    add_body(doc, "开放代码在控制领域明显少于物理模拟。权威装置 PCS、诊断映射、实时驱动、保护接口和配置通常受设施治理；开放资产更适合承担研究基线、接口验证、控制器设计、合成诊断和自动测试。报告中的代码关系标签用于防止把‘通用框架可用’误写成‘论文或装置实现已公开’。")
    rows = [
        ["动态植物/响应", "DINA、MEQ、CREATE-NL、GSevolve、FreeGSNKE、TokaMaker", "场景、自由边界、线性化响应、控制器设计"],
        ["实时状态", "EFIT/rtEFIT、LIUQE、P-EFIT、EQUINOX、RAPTOR/RAPDENS", "边界、平衡、q/剖面、预测状态与质量标志"],
        ["控制设计与仿真", "PCSSP、Simulink/Stateflow、TokSys、MATLAB/Python", "算法设计、代码生成、SIL、测试向量与证据"],
        ["实时框架", "DIII-D PCS、MARTe2、ITER RTF/CODAC、EPICS", "调度、I/O、网络、状态机、算法部署"],
        ["数据与线程", "MDSplus、IMAS/OMAS、UDA、配置数据库", "时序、语义、装置配置、版本和回放"],
        ["独立保护", "机器保护、联锁、缓解系统、核/人员安全系统", "越限处置与最终权威；不由数字孪生替代"],
    ]
    add_simple_table(doc, ["层", "代表工具/系统", "主要责任"], rows, [1800, 3450, 4110], font_size=8.2)
    figure = figures_dir / "control-pcs-layers-nature.png"
    if figure.exists():
        add_figure(doc, figure, "图 16-1  PCS、孪生与独立保护的分层接口。控制算法可以演进，责任与最终动作权限必须保持可审计。")
    doc.add_heading("建议的两层软件基线", level=2)
    add_body(doc, "开放协作层使用 PCSSP、FreeGSNKE/TokaMaker、TORAX/TRANSP、MARTe2、IMAS/OMAS 等建立公共接口、测试和高校合作环境；设施权威层连接 DINA/MEQ、实时平衡、装置 PCS、真实 I/O 与受控数据。两层共享单位、坐标、时间戳、配置、测试向量和证据格式，但开放层不冒充装置权威。")


def add_gap_and_roadmap(doc: Document, synthesis: dict[str, str], figures_dir: Path) -> None:
    add_control_kicker(doc, "17 / DISTANCE TO A CONTROL DIGITAL TWIN", page_break_before=True)
    doc.add_heading("集成控制距离数字孪生还欠缺什么", level=1)
    gaps = [
        ("配置权威", "模型、传感器、执行器、线圈连接、壁、诊断标定与限值必须绑定到具体装置/炮次和版本。"),
        ("可用的实时状态", "不仅输出数值，还要输出时间戳、质量、置信度、残差、诊断覆盖和降级状态。"),
        ("多任务契约", "目标、约束、优先级、执行器能力、未满足请求和切换原因需要机器可读。"),
        ("系统级 VVUQ", "慢模型、快速代理、控制器、PCS、网络、电源和保护接口要在同一 SIL/HIL 证据链验证。"),
        ("失败与超域证据", "训练外工况、执行器故障、缺测、时延、模型分歧和失败炮次必须成为版本验收的一部分。"),
        ("责任与权限", "候选建议、操作员批准、PCS 动作、机器保护和安全系统的权威边界必须可审计。"),
        ("持续校准", "每炮更新预测—实测残差，但生产控制模型不能未经重新验证而在线自我改变。"),
        ("电厂目标", "从等离子体性能扩展到部件寿命、RAMI、氚、热循环、净电效率、维护与安全论证。"),
    ]
    add_simple_table(doc, ["缺口", "必须补齐的能力"], [[a, b] for a, b in gaps], [1900, 7460], font_size=8.5)
    add_callout(doc, "关键差异", "集成模拟主要回答‘给定场景会怎样’，先进控制回答‘如何把状态推向目标’，数字孪生还必须回答‘当前真实装置是什么状态、模型为何可信、动作由谁批准、结果怎样反证模型，以及下一版如何受控发布’。", fill=PALE_PURPLE, accent=PURPLE)

    add_control_kicker(doc, "18 / FUSIONDIGITAL ROADMAP", page_break_before=True)
    doc.add_heading("从 DINA/MEQ 控制服务走向聚变堆与电厂", level=1)
    if "面向 FusionDigital 的路线建议" in synthesis:
        add_markdown_text(doc, synthesis["面向 FusionDigital 的路线建议"], heading_level=2)
    figure = figures_dir / "control-digital-twin-roadmap-nature.png"
    if figure.exists():
        add_figure(doc, figure, "图 18-1  FusionDigital 控制能力路线：可信回放、数字影子、跨任务预测、SIL/HIL、有限闭环和电厂级持续运行。每一级以证据门而非接入代码数量验收。")
    rows = [
        ["C0", "磁控制可信回放", "DINA/MEQ 资产包、合成磁诊断、真实控制器、线圈/电源约束、历史炮次基准", "重放可重复；模型/配置/误差可追溯"],
        ["C1", "控制数字影子", "实时/准实时状态、候选动作与风险预测，不写执行器", "每炮形成预测—实测残差和 OOD 报告"],
        ["C2", "剖面/MHD/排热协同", "快速输运、模式分析、辐射前沿、热负荷与执行器能力模型", "共享执行器冲突可预测、可解释"],
        ["C3", "系统级 SIL/HIL", "PCS 实码、I/O、网络、电源或仿真器、故障注入和最坏时延", "自动化证据通过，降级与回退可复现"],
        ["C4", "有限装置闭环", "低风险参考治理、执行器分配或局部 MPC，经操作治理逐项放权", "受限适用域 E4；版本冻结并可回退"],
        ["C5", "堆/电厂控制孪生", "燃烧、排热、工程限值、氚/RAMI、热循环、维护与电网协调", "长期可用率、安全论证和生命周期治理"],
    ]
    add_simple_table(doc, ["级", "能力", "交付", "验收门"], rows, [650, 1700, 3950, 3060], font_size=7.6)


def add_research_notes(doc: Document, notes_dir: Path) -> dict[str, int]:
    add_control_kicker(doc, "19 / EXTENDED SYNTHESIS NOTES", page_break_before=True)
    doc.add_heading("专题综合调研笔记", level=1)
    add_body(doc, "以下三组综合笔记保留了按核心连续任务、稳定性/排热/功率/保护以及 PCS/装置视角形成的长篇分析，用于专家交流时追溯分类判断、成熟度边界、负面证据和迁移风险。其内容与结构化条目互相校验，但不替代原始论文。")
    seen_paragraphs = set(iter_body_paragraph_text(doc))
    stats = {"files": 0, "addedParagraphs": 0, "skippedDuplicateParagraphs": 0}
    for filename, label in (
        ("core_tasks_notes.md", "核心连续控制任务"),
        ("protection_power_tasks_notes.md", "稳定性、排热、功率与保护任务"),
        ("pcs_devices_notes.md", "PCS 框架与装置控制生态"),
    ):
        path = notes_dir / filename
        if not path.exists():
            continue
        stats["files"] += 1
        doc.add_heading(label, level=2)
        added, skipped = add_markdown_text_deduplicated(
            doc,
            path.read_text(encoding="utf-8"),
            seen_paragraphs,
            heading_level=3,
        )
        stats["addedParagraphs"] += added
        stats["skippedDuplicateParagraphs"] += skipped
    if stats["skippedDuplicateParagraphs"]:
        add_callout(
            doc,
            "扩展笔记去重",
            f"保留 {stats['files']} 组扩展笔记，新增 {stats['addedParagraphs']} 个正文段落；跳过 {stats['skippedDuplicateParagraphs']} 个与前文章节或其他笔记完全相同的正文段落。标题与章节结构仍保留。",
            fill=PALE,
            accent=DEEP,
        )
    return stats


def add_appendix(doc: Document, entries: list[dict[str, Any]], devices: list[dict[str, Any]], task_meta: dict[str, Any]) -> None:
    add_control_kicker(doc, "APPENDICES / INDEX", page_break_before=True)
    doc.add_heading("附录 A：工作—任务—装置—证据索引", level=1)
    rows = []
    for item in entries:
        rows.append([item["id"], item["primaryTask"], item["titleZh"], "；".join(item["devices"]), f"{item['evidenceLevel']}/{item['deploymentLevel']}"])
    add_simple_table(doc, ["ID", "任务", "工作", "装置", "证据/部署"], rows, [1100, 650, 3480, 2880, 1250], font_size=6.8)
    appendix_b = doc.add_heading("附录 B：装置快速索引", level=1)
    appendix_b.paragraph_format.page_break_before = True
    rows = [[device["name"], device["country"], ", ".join(device["primaryTasks"]), str(len(device["papers"])), str(len(device["code"]))] for device in devices]
    add_simple_table(doc, ["装置", "国家/地区", "任务", "来源", "代码/软件"], rows, [2100, 1500, 2900, 1430, 1430], font_size=7.4)
    doc.add_heading("附录 C：证据解释注意事项", level=1)
    add_body(
        doc,
        "说明：信息按公开证据截止日形成，‘未公开’不等于不存在；跨装置迁移须重新验证。数字孪生/AI 不替代保护与法定验证，动作须通过独立权限与物理约束门。",
        bold_prefix="说明：",
        after=0,
    )


def task_chapter_inputs(entries: list[dict[str, Any]]) -> dict[str, tuple[list[dict[str, Any]], list[dict[str, Any]]]]:
    """Build and validate the unique full-text and cross-reference chapter routes."""
    entry_ids = [item["id"] for item in entries]
    duplicate_ids = sorted(item_id for item_id, count in Counter(entry_ids).items() if count > 1)
    if duplicate_ids:
        raise ValueError(f"Duplicate work ids prevent unique task-chapter routing: {', '.join(duplicate_ids)}")

    chapters: dict[str, tuple[list[dict[str, Any]], list[dict[str, Any]]]] = {}
    routed_primary_ids: list[str] = []
    for task in TASK_ORDER:
        primary_items = [item for item in entries if item["primaryTask"] == task]
        related_items = [
            item
            for item in entries
            if item["primaryTask"] != task and task in item["relatedTasks"]
        ]
        if not primary_items:
            raise ValueError(f"Task chapter {task} has no primary work to render as full text")
        chapters[task] = (primary_items, related_items)
        routed_primary_ids.extend(item["id"] for item in primary_items)

    if Counter(routed_primary_ids) != Counter(entry_ids):
        missing = sorted(set(entry_ids) - set(routed_primary_ids))
        repeated = sorted(item_id for item_id, count in Counter(routed_primary_ids).items() if count > 1)
        raise ValueError(
            "Every work must render in exactly one primary-task chapter; "
            f"missing={missing or 'none'} repeated={repeated or 'none'}"
        )
    return chapters


def build_report(landscape_path: Path, devices_path: Path, notes_dir: Path, figures_dir: Path, output_path: Path) -> None:
    landscape = json.loads(landscape_path.read_text(encoding="utf-8"))
    devices_payload = json.loads(devices_path.read_text(encoding="utf-8"))
    entries = landscape["entries"]
    devices = devices_payload["devices"]
    task_meta = landscape["taskMeta"]
    synthesis = parse_markdown_sections(notes_dir / "synthesis_notes.md")
    stats = {**landscape["statistics"], "devices": len(devices)}
    as_of = landscape.get("asOf")
    if not isinstance(as_of, str) or not as_of.strip():
        raise ValueError("landscape.asOf is required so the report cover uses the evidence cutoff from the landscape")
    chapters = task_chapter_inputs(entries)

    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    configure_section(section)
    configure_control_header(section)
    doc.core_properties.title = "聚变集成控制与 PCS 技术图谱研究报告"
    doc.core_properties.subject = "控制任务、装置 PCS、论文、代码、验证与聚变数字孪生路线"
    doc.core_properties.author = "新奥聚变人工智能团队 / FusionDigital"
    doc.core_properties.keywords = "FusionDigital, tokamak control, plasma control system, PCS, digital twin, DINA, MEQ, V&V"

    add_cover(doc, figures_dir, stats, as_of)
    add_toc(doc, task_meta)
    add_executive(doc, entries, devices, task_meta, figures_dir)
    add_methodology(doc, landscape, figures_dir)
    add_evidence_method(doc, landscape, figures_dir)
    add_integrated_twin_boundary(doc, synthesis)
    used_note_headings: set[str] = set()
    for chapter, task in enumerate(TASK_ORDER, 5):
        note_heading = NOTE_TASK_HEADING[task]
        primary_items, related_items = chapters[task]
        add_task_chapter(
            doc,
            task,
            primary_items,
            related_items,
            task_meta,
            synthesis,
            figures_dir,
            chapter,
            include_synthesis=note_heading not in used_note_headings,
        )
        used_note_headings.add(note_heading)
    add_device_chapter(doc, devices, entries, synthesis)
    add_pcs_toolchain(doc, figures_dir)
    add_gap_and_roadmap(doc, synthesis, figures_dir)
    notes_stats = add_research_notes(doc, notes_dir)
    add_appendix(doc, entries, devices, task_meta)

    text_stats = body_text_statistics(doc)
    if text_stats["deduplicatedCjk"] < 50_000:
        raise ValueError(
            "Report fails the 50,000-CJK substantive-body gate after exact paragraph de-duplication: "
            f"deduplicatedCjk={text_stats['deduplicatedCjk']} rawCjk={text_stats['rawCjk']} "
            f"uniqueParagraphs={text_stats['uniqueParagraphs']} duplicateParagraphs={text_stats['duplicateParagraphs']}"
        )
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"saved={output_path}")
    print(
        f"works={len(entries)} devices={len(devices)} "
        f"deduplicated_cjk_chars={text_stats['deduplicatedCjk']} raw_cjk_chars={text_stats['rawCjk']} "
        f"deduplicated_body_chars={text_stats['deduplicatedChars']} raw_body_chars={text_stats['rawChars']} "
        f"body_paragraphs={text_stats['paragraphs']} unique_body_paragraphs={text_stats['uniqueParagraphs']} "
        f"notes_duplicate_paragraphs_skipped={notes_stats['skippedDuplicateParagraphs']} "
        f"tables={len(doc.tables)} figures={len(doc.inline_shapes)}"
    )


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--landscape", type=Path, required=True)
    parser.add_argument("--devices", type=Path, required=True)
    parser.add_argument("--notes-dir", type=Path, required=True)
    parser.add_argument("--figures-dir", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    build_report(args.landscape, args.devices, args.notes_dir, args.figures_dir, args.output)


if __name__ == "__main__":
    main()
