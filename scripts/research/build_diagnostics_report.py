#!/usr/bin/env python3
"""Build the reviewed FusionDigital fusion-diagnostics technical report."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# compact_reference_guide preset, with one named CJK font override.
PAGE_W = 8.5
PAGE_H = 11.0
MARGIN = 1.0
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
FONT_ASCII = "Calibri"
FONT_CJK = "Microsoft YaHei"
INK = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PALE_BLUE = "E8EEF5"
PALE_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
ORANGE = "F47C20"
CYAN = "008C95"
PURPLE = "6D5BD0"
MUTED = "566573"
WHITE = "FFFFFF"
LINE = "C9D3DD"

TASK_ORDER = [f"DG{i}" for i in range(12)]


def set_font(run, *, size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = FONT_ASCII
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_ASCII)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_ASCII)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_CJK)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.keep_together = True


def keep_row_with_next(row) -> None:
    """Keep a table row with the row that follows it.

    Word does not expose a row-level ``keepNext`` switch, so applying the
    paragraph property to every cell in the header is the most portable way
    to prevent a repeated header from being stranded at the foot of a page.
    """
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.keep_with_next = True


def remove_trailing_spacing_paragraphs(doc: Document) -> None:
    """Remove only empty layout spacer paragraphs at the document tail.

    A ``page_break_before`` paragraph following an empty spacer can leave the
    spacer alone on an otherwise blank page.  Generated page-break paragraphs
    contain a run/``w:br`` and are therefore deliberately preserved.
    """
    body = doc.element.body
    for element in reversed(list(body)):
        if element.tag == qn("w:sectPr"):
            continue
        if element.tag != qn("w:p"):
            break
        if element.findall(".//" + qn("w:r")) or element.findall(".//" + qn("w:hyperlink")):
            break
        body.remove(element)


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != CONTENT_DXA:
        raise ValueError(f"Table widths must sum to {CONTENT_DXA}, got {sum(widths)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_hyperlink(paragraph, label: str, url: str, *, color: str = CYAN, bold: bool = False) -> None:
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT_ASCII)
    r_fonts.set(qn("w:hAnsi"), FONT_ASCII)
    r_fonts.set(qn("w:eastAsia"), FONT_CJK)
    r_pr.append(r_fonts)
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    r_pr.append(color_node)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    if bold:
        r_pr.append(OxmlElement("w:b"))
    run.append(r_pr)
    node = OxmlElement("w:t")
    node.text = label
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_font(run, size=8, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_run = OxmlElement("w:r")
    field_run.extend([begin, instr, separate, value, end])
    paragraph._p.append(field_run)
    tail = paragraph.add_run(" 页")
    set_font(tail, size=8, color=MUTED)


def ensure_numbering(doc: Document) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element
    existing_abs = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_start = max(existing_abs or [0]) + 1
    num_start = max(existing_num or [0]) + 1

    def add_definition(abstract_id: int, num_id: int, fmt: str, marker: str) -> None:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), marker)
        lvl.append(lvl_text)
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        lvl.append(lvl_jc)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        p_pr.append(ind)
        lvl.append(p_pr)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)

    add_definition(abstract_start, num_start, "bullet", "•")
    add_definition(abstract_start + 1, num_start + 1, "decimal", "%1.")
    return num_start, num_start + 1


def apply_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25


def configure_styles(doc: Document) -> tuple[int, int]:
    section = doc.sections[0]
    section.page_width = Inches(PAGE_W)
    section.page_height = Inches(PAGE_H)
    section.top_margin = Inches(MARGIN)
    section.bottom_margin = Inches(MARGIN)
    section.left_margin = Inches(MARGIN)
    section.right_margin = Inches(MARGIN)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_ASCII
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_ASCII)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_ASCII)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = FONT_ASCII
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_ASCII)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_ASCII)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = FONT_ASCII
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    caption.font.size = Pt(8.5)
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = False
    return ensure_numbering(doc)


def setup_header_footer(section) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("FusionDigital  /  聚变诊断技术图谱")
    set_font(run, size=8, color=CYAN, bold=True)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), LINE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)


def add_kicker(doc: Document, value: str, *, page_break_before: bool = False) -> None:
    if page_break_before:
        # Spacer paragraphs are useful between ordinary blocks but must not
        # become the sole occupant of a page before a new chapter.
        remove_trailing_spacing_paragraphs(doc)
    p = doc.add_paragraph()
    if page_break_before:
        p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(value.upper())
    set_font(run, size=8.5, color=CYAN, bold=True)
    run.font.all_caps = True


def add_body(doc: Document, value: str, *, bold_label: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.widow_control = True
    if bold_label:
        label = p.add_run(bold_label)
        set_font(label, bold=True, color=DARK_BLUE)
    run = p.add_run(value)
    set_font(run)


def add_bullet(doc: Document, value: str, bullet_num_id: int) -> None:
    p = doc.add_paragraph()
    apply_num(p, bullet_num_id)
    run = p.add_run(value)
    set_font(run)


def add_callout(doc: Document, label: str, value: str, *, fill: str = CALLOUT, accent: str = CYAN) -> None:
    if not str(value or "").strip():
        return
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    keep_row_together(table.rows[0])
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    lr = p.add_run(label.upper())
    set_font(lr, size=8.5, color=accent, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    vr = p2.add_run(value)
    set_font(vr, size=10, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[int],
    *,
    font_size: float = 8.4,
    compact: bool = False,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        if compact:
            set_cell_margins(cell, top=55, bottom=55, start=100, end=100)
        set_cell_shading(cell, PALE_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(value)
        set_font(run, size=font_size, color=INK, bold=True)
    # Keep the header and first data row as an indivisible opening unit.  The
    # remaining rows may paginate naturally, retaining repeated headers.
    if rows:
        keep_row_with_next(header)
    for row_values in rows:
        row = table.add_row()
        keep_row_together(row)
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            if compact:
                set_cell_margins(cell, top=45, bottom=45, start=100, end=100)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05 if compact else 1.15
            cell_text = str(value)
            if cell_text.startswith("https://"):
                add_hyperlink(p, "打开原始来源 ↗", cell_text, color=CYAN)
            else:
                run = p.add_run(cell_text)
                set_font(run, size=font_size, color=INK)
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(3)


def add_figure(doc: Document, path: Path, caption: str, *, width: float = 6.0, alt: str | None = None) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", alt or caption)
    cap = doc.add_paragraph(caption, style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_metric_strip(doc: Document, metrics: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=len(metrics))
    widths = [CONTENT_DXA // len(metrics)] * len(metrics)
    widths[-1] += CONTENT_DXA - sum(widths)
    set_table_geometry(table, widths)
    for idx, (value, label) in enumerate(metrics):
        cell = table.cell(0, idx)
        set_cell_shading(cell, INK)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(value)
        set_font(run, size=15, color=ORANGE if idx % 2 == 0 else CYAN, bold=True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        label_run = p2.add_run(label)
        set_font(label_run, size=7.5, color=WHITE)


def add_cover(doc: Document, landscape: dict[str, Any], devices: list[dict[str, Any]], figures_dir: Path) -> None:
    stats = landscape["statistics"]
    p0 = doc.add_paragraph()
    # The full cover, including version/contact metadata, is budgeted for one
    # page.  Keep the hierarchy but avoid a decorative second metadata page.
    p0.paragraph_format.space_before = Pt(24)
    p0.paragraph_format.space_after = Pt(10)
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run("FUSIONDIGITAL / FUSION DIAGNOSTICS ATLAS / 2026")
    set_font(r0, size=9, color=CYAN, bold=True)
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(4)
    r1 = p1.add_run("聚变诊断技术图谱")
    set_font(r1, size=30, color=INK, bold=True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(8)
    r2 = p2.add_run("从仪器、标定与反演，到合成诊断、装置应用和数字孪生")
    set_font(r2, size=15, color=DARK_BLUE, bold=True)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(12)
    r3 = p3.add_run("Fusion Diagnostics, Synthetic Observation, Integrated Inference and Digital-Twin Roadmap")
    set_font(r3, size=10.5, color=MUTED, italic=True)
    add_callout(doc, "DOMAIN BOUNDARY", landscape["namingBoundary"], fill="EEF7F7", accent=CYAN)
    add_figure(doc, figures_dir / "diagnostics-measurement-chain-nature.png", "图 0-1  聚变诊断的测量—反演—状态—决策链。", width=5.2)
    add_metric_strip(doc, [
        (str(stats["total"]), "项唯一工作"),
        (str(stats["uniquePapers"]), "篇一手来源"),
        (str(len(devices)), "个装置档案"),
        ("DG0-DG11", "十二类诊断任务"),
    ])
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(5)
    p4.paragraph_format.space_after = Pt(0)
    r4 = p4.add_run(f"新奥聚变人工智能团队  |  研究截止 {landscape['asOf']}  |  版本 1.0  |  tianshao1992@gmail.com")
    set_font(r4, size=8, color=MUTED)
    doc.add_page_break()


def add_reading_map(doc: Document, landscape: dict[str, Any]) -> None:
    add_kicker(doc, "CONTENTS / READING MAP")
    doc.add_heading("目录与阅读路径", level=1)
    rows = [
        ["01", "执行摘要：诊断为什么是数字孪生的事实入口"],
        ["02", "科普：从物理量到传感器信号，再从信号回到状态"],
        ["03", "DG0-DG11 分类、技术族和证据/部署口径"],
        ["04", "聚变装置、电厂与全生命周期的诊断需求"],
    ]
    for index, task in enumerate(TASK_ORDER, 5):
        meta = landscape["taskMeta"][task]
        rows.append([f"{index:02d}", f"{task}  {meta['label']}"])
    rows.extend([
        ["16A", "跨类别综合判断：诊断原理、合成观测、VVUQ 与工程状态闭环"],
        ["17", "按装置反查诊断系统、代表工作、论文与代码"],
        ["18", "从集成诊断到数字孪生仍缺什么"],
        ["19", "FusionDigital 分阶段路线与联合攻关建议"],
        ["20", "论文、代码和数据索引"],
    ])
    add_table(doc, ["章节", "主题"], rows, [900, 8460], font_size=9)
    add_callout(doc, "推荐路径", "管理者先阅读第 1、4、18、19 章；诊断专家按 DG 分类进入条目；装置合作从第 17 章反查；软件团队重点阅读 DG9-DG11 与附录代码关系。")
    doc.add_page_break()


def read_synthesis(notes_dir: Path) -> dict[str, str]:
    return read_markdown_notes(notes_dir / "synthesis_notes.md")


def read_markdown_notes(path: Path) -> dict[str, str]:
    if not path.exists():
        return {}
    sections: dict[str, list[str]] = {}
    current = "导言"
    sections[current] = []
    for line in path.read_text(encoding="utf-8").splitlines():
        if line.startswith("## "):
            current = re.sub(r"^\d+\.\s*", "", line[3:].strip())
            sections.setdefault(current, [])
        elif not line.startswith("# "):
            sections[current].append(line)
    return {key: "\n".join(value).strip() for key, value in sections.items() if "\n".join(value).strip()}


def add_markdown_section(doc: Document, content: str, bullet_num_id: int) -> None:
    pending: list[str] = []
    fenced: list[str] = []
    in_fence = False

    def clean_inline(value: str) -> str:
        value = re.sub(r"\*\*(.+?)\*\*", r"\1", value)
        return re.sub(r"`([^`]+)`", r"\1", value)

    def flush() -> None:
        if pending:
            add_body(doc, " ".join(part.strip() for part in pending if part.strip()))
            pending.clear()

    for raw in content.splitlines():
        line = raw.strip()
        if line.startswith("```"):
            flush()
            if in_fence:
                if fenced:
                    add_callout(doc, "OBSERVATION MODEL", "\n".join(fenced), fill="EEF7F7")
                fenced.clear()
                in_fence = False
            else:
                in_fence = True
            continue
        if in_fence:
            if line:
                fenced.append(clean_inline(line))
            continue
        if not line:
            flush()
        elif line.startswith("### "):
            flush()
            doc.add_heading(line[4:].strip(), level=3)
        elif re.match(r"^\d+\.\s+", line):
            flush()
            add_bullet(doc, re.sub(r"^\d+\.\s+", "", line), bullet_num_id)
        elif line.startswith("- "):
            flush()
            add_bullet(doc, clean_inline(line[2:].strip()), bullet_num_id)
        elif re.match(r"^https://", line):
            flush()
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            add_hyperlink(p, line, line)
        else:
            # Markdown tables are useful in the source notes but render poorly as
            # literal pipes in Word. Preserve their information as compact rows.
            if line.startswith("|") and line.endswith("|"):
                if re.fullmatch(r"\|[\s:|-]+\|", line):
                    continue
                flush()
                cells = [clean_inline(cell.strip()) for cell in line.strip("|").split("|")]
                add_body(doc, " · ".join(cell for cell in cells if cell))
            else:
                pending.append(clean_inline(line))
    if in_fence and fenced:
        add_callout(doc, "OBSERVATION MODEL", "\n".join(fenced), fill="EEF7F7")
    flush()


def add_front_matter(doc: Document, landscape: dict[str, Any], notes: dict[str, str], figures_dir: Path, bullet_num_id: int) -> None:
    add_kicker(doc, "01 / EXECUTIVE SUMMARY")
    doc.add_heading("诊断是聚变数字孪生的事实入口，也是最容易被低估的系统工程", level=1)
    add_callout(doc, "核心判断", "数字孪生不是把所有传感器接入数据库，而是让每个状态都能回到仪器、标定、几何、算法、不确定度、配置与装置证据；合成诊断是桥梁，AI 是加速器，独立保护和安全系统保留最终权威。", fill="FFF2E7", accent=ORANGE)
    for title in list(notes)[:4]:
        doc.add_heading(title, level=2)
        add_markdown_section(doc, notes[title], bullet_num_id)
    add_figure(doc, figures_dir / "diagnostics-taxonomy-nature.png", "图 1-1  DG0-DG11 主分类与真实测量、计算反演、实时决策之间的关系。", width=6.1)

    add_kicker(doc, "02 / MEASUREMENT CHAIN", page_break_before=True)
    doc.add_heading("从物理量到可信状态：理解聚变诊断的最短路径", level=1)
    for title in list(notes)[4:8]:
        doc.add_heading(title, level=2)
        add_markdown_section(doc, notes[title], bullet_num_id)
    add_figure(doc, figures_dir / "diagnostics-timescale-nature.png", "图 2-1  聚变诊断跨越微秒事件、脉冲运行、维护周期和电厂寿命的嵌套时间尺度。", width=6.1)

    add_kicker(doc, "03 / TAXONOMY & EVIDENCE", page_break_before=True)
    doc.add_heading("十二类任务、八类技术原理与两条成熟度轴", level=1)
    task_rows = []
    for task in TASK_ORDER:
        meta = landscape["taskMeta"][task]
        task_rows.append([task, meta["label"], meta["en"], "横切" if meta["role"] == "cross-cutting" else "测量对象"])
    add_table(doc, ["ID", "中文", "English", "角色"], task_rows, [700, 3100, 4560, 1000], font_size=7.8)
    add_body(doc, "八类技术族用于按仪器原理检索：磁与感应、微波/毫米波、激光与散射、光学/谱学/成像、核与粒子、探针/采样、工程传感以及计算反演。一个工作可有多个技术标签，但只能有一个主 DG 分类。")
    evidence_rows = [[key, value, key.replace("E", "D") if key != "E0" else "—", landscape["deploymentScale"].get(key.replace("E", "D"), "—")] for key, value in landscape["evidenceScale"].items()]
    add_table(doc, ["证据", "证据含义", "部署", "部署含义"], evidence_rows, [850, 3650, 850, 4010], font_size=8.2)
    add_callout(doc, "禁止外推", "E4 只说明真实装置在线、实时或常规使用，不自动等于 D5。D5 必须有正式审批、配置责任、独立测试、质量保证和生命周期证据。")

    add_kicker(doc, "04 / PLANT & LIFECYCLE", page_break_before=True)
    doc.add_heading("从实验装置到聚变电厂：诊断目标从丰富观测转向长期可信与可维护", level=1)
    for title in list(notes)[8:12]:
        doc.add_heading(title, level=2)
        add_markdown_section(doc, notes[title], bullet_num_id)
    add_figure(doc, figures_dir / "diagnostics-digital-twin-architecture-nature.png", "图 4-1  面向数字孪生的诊断参考架构：实体、时间/配置、观测、模型、决策与证据治理。", width=6.1)
    source_titles = [title for title in notes if "关键一手来源" in title]
    if source_titles:
        doc.add_heading("研究基线与一手来源边界", level=2)
        add_markdown_section(doc, notes[source_titles[0]], bullet_num_id)


def join_values(values: Iterable[str]) -> str:
    return "；".join(value for value in values if value) or "未单独说明"


def add_work(doc: Document, work: dict[str, Any], sequence: int) -> None:
    heading = doc.add_heading(f"{work['primaryTask']}.{sequence:02d}  {work['title']}", level=3)
    heading.paragraph_format.keep_with_next = True
    if work.get("titleEn"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(work["titleEn"])
        set_font(run, size=9, color=MUTED, italic=True)
    meta = [[work["id"], work["evidenceLevel"], work["deploymentLevel"], join_values(work["techniqueFamilies"]), join_values(item["name"] for item in work["devices"])]]
    add_table(doc, ["工作 ID", "证据", "部署", "技术族", "装置"], meta, [1200, 750, 750, 2100, 4560], font_size=7.8)
    add_body(doc, work["problem"], bold_label="解决问题：")
    add_body(doc, work["measurementPrinciple"], bold_label="测量原理：")
    add_body(doc, join_values(work["quantities"]), bold_label="测量产品：")
    add_body(doc, f"时间：{work['temporalScale']}；空间：{work['spatialScale']}", bold_label="时空分辨与覆盖：")
    add_body(doc, join_values(work["hardware"]), bold_label="仪器与传感器：")
    add_body(doc, work["calibration"], bold_label="标定与质量：")
    add_body(doc, work["inference"], bold_label="反演与分析：")
    add_body(doc, work["validation"], bold_label="验证与应用：")
    add_body(doc, work["limitations"], bold_label="局限与边界：")
    add_callout(doc, "DIGITAL-TWIN RELEVANCE", work["twinRelevance"], fill="EEF7F7", accent=CYAN)
    source_rows = []
    for paper in work["papers"]:
        source_rows.append([str(paper["year"]), paper["title"], paper["venue"], paper["sourceType"], paper["url"] or "未提供"])
    for artifact in work["code"]:
        source_rows.append(["软件", artifact["name"], artifact["status"], artifact["relation"], artifact["url"] or "未公开"])
    add_table(doc, ["年份/类型", "论文或软件", "期刊/开放性", "来源类型/关系", "链接"], source_rows, [850, 2500, 1600, 2300, 2110], font_size=7.2)


def add_task_chapters(doc: Document, landscape: dict[str, Any], figures_dir: Path) -> None:
    works = landscape["entries"]
    chapter_no = 5
    for task in TASK_ORDER:
        meta = landscape["taskMeta"][task]
        primary = [item for item in works if item["primaryTask"] == task]
        related = [item for item in works if task in item["relatedTasks"] and item["primaryTask"] != task]
        add_kicker(doc, f"{chapter_no:02d} / {task} / {meta['en']}", page_break_before=True)
        doc.add_heading(meta["label"], level=1)
        add_callout(doc, "CHAPTER SCOPE", f"本章全文展开 {len(primary)} 项主工作，并列出 {len(related)} 项跨域关联工作。关联工作只在其主任务章节全文展开，避免重复计数。", fill=PALE_GRAY)
        figure_name = {
            "DG0": "diagnostics-measurement-chain-nature.png",
            "DG9": "diagnostics-synthetic-loop-nature.png",
            "DG10": "diagnostics-inference-graph-nature.png",
            "DG11": "diagnostics-realtime-governance-nature.png",
        }.get(task)
        if figure_name:
            add_figure(doc, figures_dir / figure_name, f"图 {chapter_no}-1  {meta['label']}的关键链路。", width=6.0)
        for index, work in enumerate(primary, 1):
            add_work(doc, work, index)
        doc.add_heading("关联工作速查", level=2)
        rows = [[item["id"], item["primaryTask"], item["title"], item["evidenceLevel"], join_values(device["name"] for device in item["devices"])] for item in related]
        if not rows:
            rows = [["—", "—", "本版没有单独标注的关联工作", "—", "—"]]
        add_table(
            doc,
            ["ID", "主任务", "标题", "证据", "装置"],
            rows,
            [1000, 900, 3760, 800, 2900],
            font_size=7.3,
            compact=True,
        )
        chapter_no += 1


def add_cross_cutting_synthesis(doc: Document, notes_dir: Path, bullet_num_id: int) -> None:
    """Add reviewed synthesis that is not reducible to individual catalogue entries."""
    sources = [
        (
            "主等离子体诊断：分类边界、装置交叉验证与证据风险",
            notes_dir / "plasma_tasks_notes.md",
            ("分类边界", "Technique family", "证据与部署", "建议装置交叉", "主要证据风险", "下一轮精化"),
        ),
        (
            "合成诊断、集成反演与实时智能：从前向模型到受治理状态",
            notes_dir / "synthetic_inference_tasks_notes.md",
            ("范围、口径", "为什么", "四类资产", "数据契约", "VVUQ", "差距矩阵", "实施路线", "代码关系"),
        ),
        (
            "装置与电厂状态诊断：从实验测量到可维护资产",
            notes_dir / "devices_plant_notes.md",
            ("结论先行", "分类与装置", "装置组别", "八条主线", "最小闭环", "证据边界"),
        ),
    ]
    add_kicker(doc, "16A / CROSS-CUTTING SYNTHESIS", page_break_before=True)
    doc.add_heading("跨类别综合判断：真正的诊断孪生必须同时闭合仪器、状态、模型与证据", level=1)
    add_callout(
        doc,
        "为何单独成章",
        "条目目录回答‘有哪些工作’，本章回答‘这些工作如何共同组成可验证的诊断系统’。内容保留跨诊断共性、装置差异、VVUQ 和工程实施判断，不用重复条目数量来充当成熟度。",
        fill="EEF7F7",
        accent=CYAN,
    )
    for group_title, path, selectors in sources:
        sections = read_markdown_notes(path)
        selected = [(title, content) for title, content in sections.items() if any(token in title for token in selectors)]
        if not selected:
            continue
        doc.add_heading(group_title, level=2)
        for title, content in selected:
            doc.add_heading(title, level=3)
            add_markdown_section(doc, content, bullet_num_id)


def add_device_chapter(doc: Document, devices: list[dict[str, Any]], works: list[dict[str, Any]], figures_dir: Path) -> None:
    by_id = {item["id"]: item for item in works}
    add_kicker(doc, "17 / DEVICE-CENTRIC INDEX", page_break_before=True)
    doc.add_heading("按装置反查：主要诊断、实时接口、数据平台、论文与代码", level=1)
    add_body(doc, "装置档案是公开证据索引，不代表所有诊断在每个实验期同时可用。建设中装置按设计/采购/安装/调试证据描述；公开资料不足的 EXL-50U、EHL-2 等装置明确保留缺口，不从相邻装置外推。")
    add_figure(doc, figures_dir / "diagnostics-device-coverage-nature.png", "图 17-1  装置与 DG0-DG11 诊断任务的公开证据覆盖矩阵。", width=6.2)
    for index, device in enumerate(devices, 1):
        doc.add_heading(f"17.{index:02d}  {device['name']}", level=2)
        meta_rows = [[device["type"], device["countryOrRegion"], device["operator"], device["status"], join_values(device["primaryTasks"])]]
        add_table(doc, ["类型", "国家/地区", "运营方", "状态", "覆盖任务"], meta_rows, [1100, 1100, 2400, 1900, 2860], font_size=7.5)
        add_body(doc, device["diagnosticSummary"], bold_label="诊断概况：")
        add_body(doc, join_values(device["diagnosticSystems"]), bold_label="主要系统：")
        add_body(doc, join_values(device["sensors"]), bold_label="测量/传感：")
        add_body(doc, join_values(device["realTimeInterfaces"]), bold_label="实时与 PCS 接口：")
        add_body(doc, join_values(device["dataPlatform"]), bold_label="数据平台：")
        if device.get("representativeWorkSummaries"):
            add_body(doc, join_values(device["representativeWorkSummaries"]), bold_label="代表工作概述：")
        add_body(doc, join_values(device["limitations"]), bold_label="证据边界：")
        work_rows = []
        for work_id in device["representativeWorks"]:
            work = by_id.get(work_id)
            if work:
                work_rows.append([work_id, work["primaryTask"], work["title"], work["evidenceLevel"], work["deploymentLevel"]])
        add_table(doc, ["工作 ID", "任务", "代表工作", "证据", "部署"], work_rows or [["—", "—", "暂无可关联的结构化工作", "—", "—"]], [1200, 750, 5300, 950, 1160], font_size=7.8)
        source_rows = [[str(p["year"]), p["title"], p["venue"], p["url"] or "未提供"] for p in device["papers"]]
        source_rows += [["软件", c["name"], f"{c['status']} / {c['relation']}", c["url"] or "未公开"] for c in device["code"]]
        add_table(doc, ["年份/类型", "论文或软件", "来源/关系", "链接"], source_rows, [900, 3200, 2760, 2500], font_size=7.3)


def add_gap_and_roadmap(doc: Document, notes: dict[str, str], figures_dir: Path, bullet_num_id: int) -> None:
    add_kicker(doc, "18 / DIGITAL-TWIN GAP", page_break_before=True)
    doc.add_heading("诊断集成距离数字孪生还缺什么", level=1)
    add_callout(doc, "判断口径", "已有仪器、数据库、合成诊断或多诊断反演，并不自动构成数字孪生。最终目标还要求实体状态同步、双向模型—观测闭环、配置权威、不确定度、实时质量、生命周期证据和受治理决策接口。", fill="FFF2E7", accent=ORANGE)
    gaps = [
        ["配置权威", "几何、坐标、通道、标定、材料和算法必须绑定到装置/炮次/维护版本。"],
        ["不确定度", "从原始噪声、响应、反演、模型偏差到融合状态的预算仍不完整。"],
        ["可观测性", "电厂端口受限与诊断失效下，关键状态的冗余和独立性尚需设计。"],
        ["前向闭环", "物理/工程模型到仪器信号、再到实测残差和模型更新的闭环尚未普遍工程化。"],
        ["实时确定性", "平均延迟不足以证明 PCS/保护接口；需最坏时延、丢包、重启和故障注入。"],
        ["跨装置迁移", "几何、噪声、场景和标签漂移使算法与 AI 不能直接跨装置复制。"],
        ["生命周期", "建设、调试、运行、维护和退役数据尚未进入统一设备树与证据主线。"],
        ["责任边界", "诊断孪生、操作员、PCS、机器保护和核安全链的权限、审批与回滚需明确。"],
    ]
    add_table(doc, ["差距", "需要补齐的能力"], gaps, [1800, 7560], font_size=8.5)

    add_kicker(doc, "19 / FUSIONDIGITAL ROADMAP", page_break_before=True)
    doc.add_heading("从磁诊断可信回放，到整厂诊断数字孪生", level=1)
    roadmap_titles = [title for title in notes if any(token in title for token in ("近期路线", "联合攻关"))]
    for title in roadmap_titles:
        doc.add_heading(title, level=2)
        add_markdown_section(doc, notes[title], bullet_num_id)
    add_figure(doc, figures_dir / "diagnostics-roadmap-nature.png", "图 19-1  FusionDigital 聚变诊断数字孪生的分阶段发展路线。", width=6.1)


def add_indices(doc: Document, works: list[dict[str, Any]]) -> None:
    add_kicker(doc, "20 / PAPER, CODE & DATA INDEX", page_break_before=True)
    doc.add_heading("论文、代码与数据索引", level=1)
    add_body(doc, "所有链接均指向论文 DOI/机构原始页面、官方文档或代码仓库。not-public 表示没有找到可公开的对应实现，其链接保持为空；enabling 表示支撑诊断工作流，但不是该论文的直接实现。")
    paper_rows = []
    seen = set()
    for work in works:
        for paper in work["papers"]:
            key = paper.get("doi") or paper.get("url") or paper["title"]
            if key in seen:
                continue
            seen.add(key)
            paper_rows.append([str(paper["year"]), paper["authors"], paper["title"], paper["venue"], paper["url"] or "未提供"])
    paper_rows.sort(key=lambda row: (row[0], row[2]), reverse=True)
    doc.add_heading("20.1  唯一论文与官方来源", level=2)
    add_table(doc, ["年", "作者", "题名", "来源", "链接"], paper_rows, [600, 1600, 3200, 1700, 2260], font_size=6.9)
    code_rows = []
    seen_code = set()
    for work in works:
        for artifact in work["code"]:
            key = (artifact["name"], artifact["url"], artifact["relation"])
            if key in seen_code:
                continue
            seen_code.add(key)
            code_rows.append([artifact["name"], artifact["status"], artifact["artifactType"], artifact["access"], artifact["relation"], artifact["url"] or "未公开"])
    doc.add_heading("20.2  代码、软件与使能框架", level=2)
    add_table(doc, ["名称", "状态", "制品", "访问", "与工作的关系", "链接"], code_rows, [1400, 1250, 1050, 900, 2600, 2160], font_size=6.9)


def body_text_stats(doc: Document) -> dict[str, int]:
    values = []
    for paragraph in doc.element.body.iter(qn("w:p")):
        value = re.sub(r"\s+", " ", "".join(node.text or "" for node in paragraph.iter(qn("w:t")))).strip()
        if value:
            values.append(value)
    unique = list(dict.fromkeys(values))
    count = lambda seq: sum(len(re.findall(r"[\u3400-\u4dbf\u4e00-\u9fff]", item)) for item in seq)
    return {"paragraphs": len(values), "uniqueParagraphs": len(unique), "rawCjk": count(values), "deduplicatedCjk": count(unique)}


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--landscape", type=Path, required=True)
    parser.add_argument("--devices", type=Path, required=True)
    parser.add_argument("--notes-dir", type=Path, required=True)
    parser.add_argument("--figures-dir", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    landscape = json.loads(args.landscape.read_text(encoding="utf-8"))
    device_payload = json.loads(args.devices.read_text(encoding="utf-8"))
    if not landscape.get("asOf"):
        raise ValueError("landscape.asOf is required")
    works = landscape["entries"]
    devices = device_payload["devices"]
    notes = read_synthesis(args.notes_dir)
    doc = Document()
    bullet_num_id, _ = configure_styles(doc)
    setup_header_footer(doc.sections[0])
    add_cover(doc, landscape, devices, args.figures_dir)
    add_reading_map(doc, landscape)
    add_front_matter(doc, landscape, notes, args.figures_dir, bullet_num_id)
    add_task_chapters(doc, landscape, args.figures_dir)
    add_cross_cutting_synthesis(doc, args.notes_dir, bullet_num_id)
    add_device_chapter(doc, devices, works, args.figures_dir)
    add_gap_and_roadmap(doc, notes, args.figures_dir, bullet_num_id)
    add_indices(doc, works)
    stats = body_text_stats(doc)
    if stats["deduplicatedCjk"] < 50_000:
        raise ValueError(f"Report substantive CJK text is below 50,000: {stats}")
    doc.core_properties.title = "聚变诊断技术图谱研究报告"
    doc.core_properties.subject = "Fusion diagnostics, synthetic diagnostics, device applications and digital-twin roadmap"
    doc.core_properties.author = "新奥聚变人工智能团队"
    doc.core_properties.comments = f"FusionDigital diagnostics atlas; as of {landscape['asOf']}"
    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)
    print(json.dumps({"output": str(args.output), "works": len(works), "devices": len(devices), **stats}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
